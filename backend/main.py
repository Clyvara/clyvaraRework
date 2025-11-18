from fastapi import FastAPI, Request, Depends, HTTPException, UploadFile, File, Header, Body
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, Dict, Any, List
from uuid import uuid4
import random
import json
import boto3
import PyPDF2
import io
from docx import Document

from openai import OpenAI
from dotenv import load_dotenv
import os
import jwt

from sqlalchemy.orm import Session
from sqlalchemy import func, or_
from database import get_db, test_connection, ChatMessage, UserInteraction, UserSession, Material, VectorIndexEntry, CarePlan, Profile, LearningPlan, LearningPlanProgress
from material_cache import (
    get_cached_text, cache_text, invalidate_cache, preload_system_materials, get_cache_stats,
    get_cached_vector_entries, cache_vector_entries, invalidate_vector_cache
)

load_dotenv()

# Supabase JWT secret for token verification
SUPABASE_JWT_SECRET = os.getenv("SUPABASE_JWT_SECRET")

# System materials user ID - materials with this user_id are accessible to all users
SYSTEM_USER_ID = "SYSTEM"

# Cache for general query embedding (since it's always the same)
_general_query_embedding_cache = None

def get_current_user(authorization: str = Header(None)):
    """Extract user info from Supabase JWT token"""
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid authorization header")
    
    token = authorization.split(" ")[1]
    
    try:
        # For Supabase tokens, decode without signature verification for now
        # This is a simplified approach for development
        payload = jwt.decode(token, options={"verify_signature": False})
        
        return {
            "user_id": payload.get("sub"),
            "email": payload.get("email"),
            "role": payload.get("role", "authenticated")
        }
    except jwt.InvalidTokenError as e:
        raise HTTPException(status_code=401, detail=f"Invalid token: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"Token error: {str(e)}")

# Initialize OpenAI client only if API key is available
openai_api_key = os.getenv("OPENAI_API_KEY")
if openai_api_key and openai_api_key != "your_openai_api_key_here":
    client = OpenAI(api_key=openai_api_key)
else:
    client = None

# AWS S3 Configuration
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
AWS_REGION = os.getenv("AWS_REGION", "us-east-2")
S3_BUCKET_NAME = os.getenv("S3_BUCKET_NAME", "clyvara-uploads")

# Initialize S3 client
if AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY:
    s3_client = boto3.client(
        's3',
        aws_access_key_id=AWS_ACCESS_KEY_ID,
        aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
        region_name=AWS_REGION
    )
else:
    s3_client = None
        
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8001", "http://localhost:8002", "http://localhost:8003"],
    allow_methods=["*"],
    allow_headers=["*"],
    allow_credentials=True,
)

@app.on_event("startup")
async def startup_event():
    """Preload system materials into cache on startup"""
    try:
        from database import get_session_local
        db = get_session_local()()
        try:
            preload_system_materials(db, SYSTEM_USER_ID)
            print("Material cache initialized")
        finally:
            db.close()
    except Exception as e:
        print(f"Warning: Could not preload system materials: {e}")

# File processing functions
def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting PDF text: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting DOCX text: {str(e)}")

def extract_text_from_file(file_content: bytes, file_type: str) -> str:
    """Extract text based on file type"""
    if file_type.lower() == "pdf":
        return extract_text_from_pdf(file_content)
    elif file_type.lower() in ["docx", "doc"]:
        return extract_text_from_docx(file_content)
    elif file_type.lower() == "txt":
        return file_content.decode('utf-8')
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")

def generate_embeddings(text: str) -> List[float]:
    """Generate embeddings for text using OpenAI"""
    if not client:
        raise HTTPException(status_code=503, detail="OpenAI client not configured")
    
    try:
        response = client.embeddings.create(
            model="text-embedding-3-small",
            input=text
        )
        return response.data[0].embedding
    except Exception as e:
        error_msg = str(e)
        # Provide more specific error messages for common issues
        if "401" in error_msg or "invalid_api_key" in error_msg.lower():
            raise HTTPException(
                status_code=500, 
                detail="Error generating embeddings: Invalid OpenAI API key. Please check your OPENAI_API_KEY in the .env file."
            )
        elif "403" in error_msg or "forbidden" in error_msg.lower():
            raise HTTPException(
                status_code=500,
                detail="Error generating embeddings: API key access denied (403). This may be due to insufficient quota, expired key, or restricted permissions. Please check your OpenAI account."
            )
        elif "429" in error_msg or "rate limit" in error_msg.lower():
            raise HTTPException(
                status_code=500,
                detail="Error generating embeddings: Rate limit exceeded. Please wait a moment and try again."
            )
        else:
            raise HTTPException(status_code=500, detail=f"Error generating embeddings: {error_msg}")

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks for better retrieval"""
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        chunks.append(chunk)
    
    return chunks

@app.get("/")
def root():
    return {"message": "Clyvara Backend API", "status": "running"}

@app.get("/health")
def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "database": "connected" if test_connection() else "disconnected"
    }

@app.get("/api/cache/stats")
def get_cache_stats_endpoint():
    """Get material cache statistics"""
    try:
        stats = get_cache_stats()
        return {
            "success": True,
            "cache_stats": stats
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

@app.get("/api/tables")
def list_tables(db: Session = Depends(get_db)):
    """List all tables in the database"""
    try:
        from sqlalchemy import text
        result = db.execute(text("""
            SELECT table_name, table_schema 
            FROM information_schema.tables 
            WHERE table_schema NOT IN ('information_schema', 'pg_catalog')
            ORDER BY table_schema, table_name
        """))
        tables = [{"schema": row[1], "table": row[0]} for row in result]
        return {"tables": tables, "count": len(tables)}
    except Exception as e:
        return {"error": f"Database error: {str(e)}"}

@app.post("/api/ensure-tables")
def ensure_tables():
    """Ensure all tables exist, create if they don't"""
    try:
        from database import init_db, get_engine
        from sqlalchemy import inspect
        # Use init_db to create schemas and tables
        success = init_db()
        if success:
            # Verify profiles table exists
            engine = get_engine()
            inspector = inspect(engine)
            tables_in_main = inspector.get_table_names(schema='main')
            profiles_exists = 'profiles' in tables_in_main
            
            return {
                "success": True, 
                "message": "All tables ensured/created successfully",
                "profiles_table_exists": profiles_exists,
                "tables_in_main_schema": tables_in_main
            }
        else:
            return {"success": False, "message": "Failed to create tables"}
    except Exception as e:
        import traceback
        return {"success": False, "error": str(e), "traceback": traceback.format_exc()}

@app.post("/api/test-auth")
def test_authentication_flow(
    test_user_id: str = "550e8400-e29b-41d4-a716-446655440000",
    test_email: str = "test@example.com",
    db: Session = Depends(get_db)
):
    """Test the complete Supabase → AWS flow"""
    try:
        # First create a user session (required for foreign key)
        test_session = UserSession(
            session_id="test-session-123",
            user_id=test_user_id,
            is_active=True
        )
        db.add(test_session)
        db.commit()
        
        # Now create a chat message with Supabase user ID
        test_message = ChatMessage(
            session_id="test-session-123",
            message_type="test",
            message_content={"test": "Supabase → AWS connection working!"},
            user_id=test_user_id,  # This would come from Supabase JWT
            response_time_ms=100,
            message_length=50
        )
        
        db.add(test_message)
        db.commit()
        db.refresh(test_message)
        
        # Verify the data was stored
        stored_message = db.query(ChatMessage).filter(
            ChatMessage.user_id == test_user_id
        ).first()
        
        return {
            "status": "success",
            "message": "Supabase → AWS connection working!",
            "supabase_user_id": test_user_id,
            "aws_stored_data": {
                "message_id": str(stored_message.id),
                "user_id": stored_message.user_id,
                "content": stored_message.message_content,
                "timestamp": stored_message.timestamp.isoformat()
            }
        }
    except Exception as e:
        return {"error": f"Connection test failed: {str(e)}"}

@app.get("/api/user-data/{user_id}")
def get_user_data(user_id: str, db: Session = Depends(get_db)):
    """Get all data for a specific Supabase user from AWS"""
    try:
        # Get chat messages for this user
        messages = db.query(ChatMessage).filter(
            ChatMessage.user_id == user_id
        ).all()
        
        # Get user interactions
        interactions = db.query(UserInteraction).filter(
            UserInteraction.user_id == user_id
        ).all()
        
        return {
            "supabase_user_id": user_id,
            "chat_messages": [
                {
                    "id": str(msg.id),
                    "content": msg.message_content,
                    "timestamp": msg.timestamp.isoformat()
                } for msg in messages
            ],
            "user_interactions": [
                {
                    "id": str(inter.id),
                    "type": inter.interaction_type,
                    "timestamp": inter.timestamp.isoformat()
                } for inter in interactions
            ],
            "total_messages": len(messages),
            "total_interactions": len(interactions)
        }
    except Exception as e:
        return {"error": f"Failed to get user data: {str(e)}"}

@app.get("/api/all-users-with-data")
def get_all_users_with_data(db: Session = Depends(get_db)):
    """Get all Supabase users who have data in AWS"""
    try:
        # Get all unique user IDs from chat messages (filter out None)
        chat_users = db.query(ChatMessage.user_id).filter(ChatMessage.user_id.isnot(None)).distinct().all()
        chat_user_ids = [str(user[0]) for user in chat_users if user[0] is not None]
        
        # Get all unique user IDs from user interactions (filter out None)
        interaction_users = db.query(UserInteraction.user_id).filter(UserInteraction.user_id.isnot(None)).distinct().all()
        interaction_user_ids = [str(user[0]) for user in interaction_users if user[0] is not None]
        
        # Combine and deduplicate
        all_user_ids = list(set(chat_user_ids + interaction_user_ids))
        
        # Get user details for each
        user_data = []
        for user_id in all_user_ids:
            messages = db.query(ChatMessage).filter(ChatMessage.user_id == user_id).count()
            interactions = db.query(UserInteraction).filter(UserInteraction.user_id == user_id).count()
            
            user_data.append({
                "supabase_user_id": user_id,
                "total_messages": messages,
                "total_interactions": interactions,
                "has_data": messages > 0 or interactions > 0
            })
        
        return {
            "users_with_data": user_data,
            "total_users": len(user_data),
            "summary": {
                "users_with_chat_messages": len([u for u in user_data if u["total_messages"] > 0]),
                "users_with_interactions": len([u for u in user_data if u["total_interactions"] > 0])
            }
        }
    except Exception as e:
        return {"error": f"Failed to get user data: {str(e)}"}

@app.get("/api/recent-chat-messages")
def get_recent_chat_messages(db: Session = Depends(get_db)):
    """Get recent chat messages from the database"""
    try:
        # Get the 10 most recent chat messages
        messages = db.query(ChatMessage).order_by(ChatMessage.timestamp.desc()).limit(10).all()
        
        return {
            "recent_messages": [
                {
                    "id": str(msg.id),
                    "user_id": str(msg.user_id),
                    "message_content": msg.message_content,
                    "timestamp": msg.timestamp.isoformat(),
                    "session_id": msg.session_id
                } for msg in messages
            ],
            "total_messages": len(messages)
        }
    except Exception as e:
        return {"error": f"Failed to get chat messages: {str(e)}"}

@app.get("/api/debug-user")
def debug_user(current_user: dict = Depends(get_current_user)):
    """Debug endpoint to check user extraction"""
    return {
        "user_info": current_user,
        "user_id": current_user.get("user_id"),
        "email": current_user.get("email")
    }

@app.post("/api/test-real-auth")
def test_real_supabase_auth(
    authorization: str = Header(None),
    db: Session = Depends(get_db)
):
    """Test with real Supabase JWT token"""
    try:
        if not authorization or not authorization.startswith("Bearer "):
            return {"error": "No valid authorization header provided"}
        
        token = authorization.split(" ")[1]
        current_user = get_current_user(token)
        
        # Create a test message with the real Supabase user
        test_message = ChatMessage(
            session_id=f"session-{current_user['sub']}",
            message_type="real_auth_test",
            message_content={"message": "Real Supabase user authenticated!", "email": current_user.get('email')},
            user_id=current_user['sub'],
            response_time_ms=50,
            message_length=100
        )
        
        db.add(test_message)
        db.commit()
        db.refresh(test_message)
        
        return {
            "status": "success",
            "message": "Real Supabase authentication working!",
            "supabase_user": {
                "id": current_user['sub'],
                "email": current_user.get('email'),
                "aud": current_user.get('aud')
            },
            "aws_stored_data": {
                "message_id": str(test_message.id),
                "user_id": test_message.user_id,
                "content": test_message.message_content,
                "timestamp": test_message.timestamp.isoformat()
            }
        }
    except Exception as e:
        return {"error": f"Real auth test failed: {str(e)}"}

@app.post("/items")
async def create_item(item: Dict):
    print("Received from POST:", item)
    return {"json data": item}

# Database API Endpoints

@app.get("/api/chat-messages")
def get_chat_messages(
    session_id: Optional[str] = None,
    user_id: Optional[str] = None,
    limit: int = 50,
    db: Session = Depends(get_db)
):
    """Get chat messages with optional filtering"""
    try:
        query = db.query(ChatMessage)
        
        if session_id:
            query = query.filter(ChatMessage.session_id == session_id)
        if user_id:
            query = query.filter(ChatMessage.user_id == user_id)
        
        messages = query.order_by(ChatMessage.timestamp.desc()).limit(limit).all()
        return {"messages": [{"id": str(msg.id), "session_id": msg.session_id, "message_type": msg.message_type, 
                             "timestamp": msg.timestamp.isoformat() if msg.timestamp else None, "user_id": str(msg.user_id) if msg.user_id else None} for msg in messages]}
    except Exception as e:
        return {"error": f"Database error: {str(e)}", "message": "Tables may not exist yet. Please create them first."}

@app.post("/api/chat-messages")
def create_chat_message(
    session_id: str,
    message_type: str,
    message_content: Dict,
    thread_id: Optional[str] = None,
    user_id: Optional[str] = None,
    response_time_ms: Optional[int] = None,
    message_length: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Create a new chat message"""
    # Use authenticated user ID from Supabase
    authenticated_user_id = current_user["user_id"]
    
    message = ChatMessage(
        session_id=session_id,
        message_type=message_type,
        message_content=message_content,
        thread_id=thread_id,
        user_id=authenticated_user_id,  # Use Supabase user ID
        response_time_ms=response_time_ms,
        message_length=message_length
    )
    
    db.add(message)
    db.commit()
    db.refresh(message)
    
    return {"message": "Chat message created", "id": message.id}

@app.get("/api/user-sessions")
def get_user_sessions(
    user_id: Optional[str] = None,
    is_active: Optional[bool] = None,
    limit: int = 50,
    db: Session = Depends(get_db)
):
    """Get user sessions with optional filtering"""
    try:
        query = db.query(UserSession)
        
        if user_id:
            query = query.filter(UserSession.user_id == user_id)
        if is_active is not None:
            query = query.filter(UserSession.is_active == is_active)
        
        sessions = query.order_by(UserSession.created_at.desc()).limit(limit).all()
        return {"sessions": [{"id": str(sess.id), "session_id": sess.session_id, "user_id": str(sess.user_id) if sess.user_id else None,
                             "created_at": sess.created_at.isoformat() if sess.created_at else None, "is_active": sess.is_active} for sess in sessions]}
    except Exception as e:
        return {"error": f"Database error: {str(e)}", "message": "Tables may not exist yet. Please create them first."}

@app.get("/api/user-interactions")
def get_user_interactions(
    user_id: Optional[str] = None,
    session_id: Optional[str] = None,
    interaction_type: Optional[str] = None,
    limit: int = 50,
    db: Session = Depends(get_db)
):
    """Get user interactions with optional filtering"""
    try:
        query = db.query(UserInteraction)
        
        if user_id:
            query = query.filter(UserInteraction.user_id == user_id)
        if session_id:
            query = query.filter(UserInteraction.session_id == session_id)
        if interaction_type:
            query = query.filter(UserInteraction.interaction_type == interaction_type)
        
        interactions = query.order_by(UserInteraction.timestamp.desc()).limit(limit).all()
        return {"interactions": [{"id": str(inter.id), "session_id": inter.session_id, "user_id": str(inter.user_id) if inter.user_id else None,
                                 "interaction_type": inter.interaction_type, "timestamp": inter.timestamp.isoformat() if inter.timestamp else None} for inter in interactions]}
    except Exception as e:
        return {"error": f"Database error: {str(e)}", "message": "Tables may not exist yet. Please create them first."}


# Chatbot Setup

with open("systemprompt.txt", "r", encoding="utf-8") as f:
    SYSTEM_PROMPT = f.read()

SESSIONS: Dict[str, Optional[str]] = {}

PRIORITY_WORDS = {
    "login","log in","signin","sign in","sign-in",
    "logout","sign out","register","sign up",
    "profile","account","menu","pricing","help","docs","dashboard"
}

def _safe_text(x: Optional[str]) -> str:
    if not isinstance(x, str):
        return ""
    return x.strip()

def summarize_context_on_server(ctx: dict, limit: int = 24) -> dict:
    """
    Keep payload tiny & focused. Prefer nav/auth items; fallback to a few elements.
    """
    if not isinstance(ctx, dict):
        return {}

    url   = _safe_text(ctx.get("url"))
    title = _safe_text(ctx.get("title"))
    headings = ctx.get("headings") or []
    elements = ctx.get("elements") or []

    clean_headings = []
    for h in headings[:8]:
        if isinstance(h, dict):
            clean_headings.append({
                "tag": _safe_text(h.get("tag")),
                "text": _safe_text(h.get("text"))[:120],
            })

    prioritized: List[dict] = []
    for el in elements:
        if not isinstance(el, dict):
            continue
        text = _safe_text(el.get("text") or el.get("ariaLabel"))
        if any(w in text.lower() for w in PRIORITY_WORDS) or _safe_text(el.get("dataQa")):
            prioritized.append({
                "tag": _safe_text(el.get("tag")),
                "text": text[:80],
                "ariaLabel": _safe_text(el.get("ariaLabel"))[:80],
                "href": _safe_text(el.get("href"))[:200],
                "id": _safe_text(el.get("id"))[:80],
                "dataQa": _safe_text(el.get("dataQa"))[:80],
                "region": _safe_text(el.get("region"))[:20],
            })
            if len(prioritized) >= limit:
                break

    if not prioritized:
        for el in elements[:limit]:
            if not isinstance(el, dict):
                continue
            prioritized.append({
                "tag": _safe_text(el.get("tag")),
                "text": _safe_text(el.get("text"))[:80],
                "ariaLabel": _safe_text(el.get("ariaLabel"))[:80],
                "href": _safe_text(el.get("href"))[:200],
                "id": _safe_text(el.get("id"))[:80],
                "dataQa": _safe_text(el.get("dataQa"))[:80],
                "region": _safe_text(el.get("region"))[:20],
            })

    return {
        "url": url,
        "title": title,
        "headings": clean_headings,
        "elements": prioritized,
    }

def build_page_context_message(ctx: Optional[dict]) -> Optional[dict]:
    """
    Convert summarized context into a compact system message the model can use.
    """
    if not ctx:
        return None
    summary = summarize_context_on_server(ctx)
    if not summary:
        return None
    content = "PAGE_CONTEXT:\n" + json.dumps(summary, ensure_ascii=False)
    return {"role": "system", "content": content}


class ChatIn(BaseModel):
    message: str
    thread_id: Optional[str] = None
    page_context: Optional[Dict[str, Any]] = None

class ChatOut(BaseModel):
    reply: str
    thread_id: str

@app.post("/chat", response_model=ChatOut)
def chat(payload: ChatIn, current_user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Chat endpoint with authentication and database storage"""
    if not client:
        raise HTTPException(status_code=503, detail="OpenAI client not configured")
    
    # Create or reuse a thread
    thread_id = payload.thread_id or str(uuid4())
    
    # Simple conversation without session management for now
    messages = [
        {"role": "system", "content": "You are a helpful AI assistant. You can answer questions about medical topics, general knowledge, current events, and provide practical information. Be helpful and informative in your responses."}
    ]
    
    messages.append({"role": "user", "content": payload.message})
    
    try:
        # Use the correct OpenAI model
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=500
        )
        
        reply = response.choices[0].message.content
        
        # Store the chat message in database with real user ID
        try:
            # First create a user session if it doesn't exist
            existing_session = db.query(UserSession).filter(UserSession.session_id == thread_id).first()
            if not existing_session:
                user_session = UserSession(
                    session_id=thread_id,
                    user_id=current_user["user_id"],
                    is_active=True
                )
                db.add(user_session)
                db.commit()
            
            # Store the chat message
            chat_message = ChatMessage(
                session_id=thread_id,
                message_type="user",
                message_content={"message": payload.message},
                user_id=current_user["user_id"],
                response_time_ms=100,  # Placeholder
                message_length=len(payload.message)
            )
            
            db.add(chat_message)
            db.commit()
            
        except Exception as db_error:
            # Log the error but don't fail the chat
            print(f"Database storage error: {db_error}")
        
        return ChatOut(reply=reply, thread_id=thread_id)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Chat error: {str(e)}")

# Care Plan API Endpoints
@app.post("/api/care-plans")
async def create_care_plan(
    care_plan_data: dict,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new care plan"""
    try:
        # Create care plan record
        care_plan = CarePlan(
            user_id=current_user['user_id'],
            title=care_plan_data.get('title', 'Untitled Care Plan'),
            patient_name=care_plan_data.get('patient_name'),
            procedure=care_plan_data.get('procedure'),
            diagnosis=care_plan_data.get('diagnosis'),
            
            # Demographics
            age=care_plan_data.get('age'),
            sex=care_plan_data.get('sex'),
            height=care_plan_data.get('height'),
            weight=care_plan_data.get('weight'),
            
            # Vital Signs
            temperature_f=care_plan_data.get('tempF'),
            blood_pressure=care_plan_data.get('bp'),
            heart_rate=care_plan_data.get('hr'),
            respiration_rate=care_plan_data.get('rr'),
            oxygen_saturation=care_plan_data.get('oxygen_saturation'),
            lmp_date=care_plan_data.get('lmp_date'),
            
            # Medical History
            past_medical_history=care_plan_data.get('pmh'),
            past_surgical_history=care_plan_data.get('psh'),
            anesthesia_history=care_plan_data.get('anesthesiaHx'),
            current_medications=care_plan_data.get('meds'),
            alcohol_use=care_plan_data.get('alcohol'),
            substance_use=care_plan_data.get('substance'),
            allergies=care_plan_data.get('allergies'),
            
            # Physical Assessment
            neurological_findings=care_plan_data.get('neuro'),
            heent_findings=care_plan_data.get('heent'),
            respiratory_findings=care_plan_data.get('resp'),
            cardiovascular_findings=care_plan_data.get('cardio'),
            gastrointestinal_findings=care_plan_data.get('gi'),
            genitourinary_findings=care_plan_data.get('gu'),
            endocrine_findings=care_plan_data.get('endo'),
            other_findings=care_plan_data.get('otherFindings'),
            
            # Airway Assessment
            mallampati_class=care_plan_data.get('mallampati'),
            ulbt_grade=care_plan_data.get('ulbt'),
            thyromental_distance=care_plan_data.get('thyromental'),
            interincisor_distance=care_plan_data.get('interincisor'),
            dentition=care_plan_data.get('dentition'),
            neck_assessment=care_plan_data.get('neck'),
            oral_mucosa=care_plan_data.get('oralMucosa'),
            
            # Laboratory Values
            sodium=care_plan_data.get('na'),
            potassium=care_plan_data.get('k'),
            chloride=care_plan_data.get('cl'),
            co2=care_plan_data.get('co2'),
            bun=care_plan_data.get('bun'),
            creatinine=care_plan_data.get('cr'),
            glucose=care_plan_data.get('glu'),
            wbc=care_plan_data.get('wbc'),
            hemoglobin=care_plan_data.get('hgb'),
            hematocrit=care_plan_data.get('hct'),
            platelets=care_plan_data.get('plt'),
            pt=care_plan_data.get('pt'),
            ptt=care_plan_data.get('ptt'),
            inr=care_plan_data.get('inr'),
            abg=care_plan_data.get('abg'),
            other_labs=care_plan_data.get('otherLabs'),
            
            # Imaging/Diagnostic Tests
            ekg=care_plan_data.get('ekg'),
            chest_xray=care_plan_data.get('cxr'),
            echocardiogram=care_plan_data.get('echo'),
            other_imaging=care_plan_data.get('otherImaging'),
            
            # Cultural/Religious Considerations
            cultural_religious_attributes=care_plan_data.get('cultural_religious_attributes'),
            
            # Generate export text for RAG indexing
            exported_text=build_care_plan_text(care_plan_data),
            export_hash=hash(str(care_plan_data))
        )
        
        db.add(care_plan)
        db.commit()
        db.refresh(care_plan)
        
        # Index the care plan content for RAG
        await index_care_plan_for_rag(care_plan, db)
        
        return {
            "success": True,
            "care_plan_id": care_plan.id,
            "message": "Care plan created successfully"
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error creating care plan: {str(e)}")

@app.get("/api/care-plans")
async def get_care_plans(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all care plans for the current user"""
    try:
        care_plans = db.query(CarePlan).filter(
            CarePlan.user_id == current_user['user_id']
        ).order_by(CarePlan.created_at.desc()).all()
        
        return {
            "success": True,
            "care_plans": [
                {
                    "id": cp.id,
                    "title": cp.title,
                    "patient_name": cp.patient_name,
                    "procedure": cp.procedure,
                    "diagnosis": cp.diagnosis,
                    "status": cp.status,
                    "created_at": cp.created_at.isoformat() if cp.created_at else None,
                    "updated_at": cp.updated_at.isoformat() if cp.updated_at else None
                }
                for cp in care_plans
            ]
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching care plans: {str(e)}")

@app.get("/api/care-plans/{care_plan_id}")
async def get_care_plan(
    care_plan_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a specific care plan by ID"""
    try:
        care_plan = db.query(CarePlan).filter(
            CarePlan.id == care_plan_id,
            CarePlan.user_id == current_user['user_id']
        ).first()
        
        if not care_plan:
            raise HTTPException(status_code=404, detail="Care plan not found")
        
        # Update last accessed
        care_plan.last_accessed = func.now()
        db.commit()
        
        return {
            "success": True,
            "care_plan": {
                "id": care_plan.id,
                "title": care_plan.title,
                "patient_name": care_plan.patient_name,
                "procedure": care_plan.procedure,
                "diagnosis": care_plan.diagnosis,
                
                # Demographics
                "age": care_plan.age,
                "sex": care_plan.sex,
                "height": care_plan.height,
                "weight": care_plan.weight,
                
                # Vital Signs
                "tempF": care_plan.temperature_f,
                "bp": care_plan.blood_pressure,
                "hr": care_plan.heart_rate,
                "rr": care_plan.respiration_rate,
                "oxygen_saturation": care_plan.oxygen_saturation,
                "lmp_date": care_plan.lmp_date,
                
                # Medical History
                "pmh": care_plan.past_medical_history,
                "psh": care_plan.past_surgical_history,
                "anesthesiaHx": care_plan.anesthesia_history,
                "meds": care_plan.current_medications,
                "alcohol": care_plan.alcohol_use,
                "substance": care_plan.substance_use,
                "allergies": care_plan.allergies,
                
                # Physical Assessment
                "neuro": care_plan.neurological_findings,
                "heent": care_plan.heent_findings,
                "resp": care_plan.respiratory_findings,
                "cardio": care_plan.cardiovascular_findings,
                "gi": care_plan.gastrointestinal_findings,
                "gu": care_plan.genitourinary_findings,
                "endo": care_plan.endocrine_findings,
                "otherFindings": care_plan.other_findings,
                
                # Airway Assessment
                "mallampati": care_plan.mallampati_class,
                "ulbt": care_plan.ulbt_grade,
                "thyromental": care_plan.thyromental_distance,
                "interincisor": care_plan.interincisor_distance,
                "dentition": care_plan.dentition,
                "neck": care_plan.neck_assessment,
                "oralMucosa": care_plan.oral_mucosa,
                
                # Laboratory Values
                "na": care_plan.sodium,
                "k": care_plan.potassium,
                "cl": care_plan.chloride,
                "co2": care_plan.co2,
                "bun": care_plan.bun,
                "cr": care_plan.creatinine,
                "glu": care_plan.glucose,
                "wbc": care_plan.wbc,
                "hgb": care_plan.hemoglobin,
                "hct": care_plan.hematocrit,
                "plt": care_plan.platelets,
                "pt": care_plan.pt,
                "ptt": care_plan.ptt,
                "inr": care_plan.inr,
                "abg": care_plan.abg,
                "otherLabs": care_plan.other_labs,
                
                # Imaging/Diagnostic Tests
                "ekg": care_plan.ekg,
                "cxr": care_plan.chest_xray,
                "echo": care_plan.echocardiogram,
                "otherImaging": care_plan.other_imaging,
                
                # Cultural/Religious Considerations
                "cultural_religious_attributes": care_plan.cultural_religious_attributes,
                
                # AI-Generated Content
                "ai_recommendations": care_plan.ai_recommendations,
                "risk_assessment": care_plan.risk_assessment,
                "monitoring_plan": care_plan.monitoring_plan,
                "medication_plan": care_plan.medication_plan,
                
                # Metadata
                "status": care_plan.status,
                "version": care_plan.version,
                "created_at": care_plan.created_at.isoformat() if care_plan.created_at else None,
                "updated_at": care_plan.updated_at.isoformat() if care_plan.updated_at else None
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching care plan: {str(e)}")

@app.post("/api/care-plans/{care_plan_id}/generate-ai")
async def generate_ai_recommendations(
    care_plan_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Generate AI recommendations for a care plan using RAG"""
    try:
        care_plan = db.query(CarePlan).filter(
            CarePlan.id == care_plan_id,
            CarePlan.user_id == current_user['user_id']
        ).first()
        
        if not care_plan:
            raise HTTPException(status_code=404, detail="Care plan not found")
        
        if not client:
            raise HTTPException(status_code=503, detail="OpenAI client not configured")
        
        # Build context for AI generation
        context = build_care_plan_context(care_plan, db)
        
        # Generate AI recommendations
        recommendations = await generate_care_plan_recommendations(context, client)
        
        # Update care plan with AI recommendations
        care_plan.ai_recommendations = recommendations.get('anesthesia_plan', '')
        care_plan.risk_assessment = recommendations.get('risk_assessment', '')
        care_plan.monitoring_plan = recommendations.get('monitoring_plan', '')
        care_plan.medication_plan = recommendations.get('medication_plan', '')
        care_plan.rag_context = context
        care_plan.rag_sources = recommendations.get('sources', [])
        care_plan.rag_confidence_score = recommendations.get('confidence_score', 0.0)
        care_plan.updated_at = func.now()
        
        db.commit()
        
        return {
            "success": True,
            "recommendations": recommendations,
            "message": "AI recommendations generated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating AI recommendations: {str(e)}")

@app.delete("/api/care-plans/{care_plan_id}")
async def delete_care_plan(
    care_plan_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a care plan"""
    try:
        care_plan = db.query(CarePlan).filter(
            CarePlan.id == care_plan_id,
            CarePlan.user_id == current_user['user_id']
        ).first()
        
        if not care_plan:
            raise HTTPException(status_code=404, detail="Care plan not found")
        
        # Delete associated RAG entries
        db.query(VectorIndexEntry).filter(
            VectorIndexEntry.source_id == care_plan_id,
            VectorIndexEntry.source_type == "care_plan"
        ).delete()
        
        # Delete care plan
        db.delete(care_plan)
        db.commit()
        
        return {
            "success": True,
            "message": "Care plan deleted successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error deleting care plan: {str(e)}")

# Helper functions for care plan operations
def build_care_plan_text(care_plan_data: dict) -> str:
    """Build text representation of care plan for RAG indexing"""
    sections = []
    
    # Patient Information
    if care_plan_data.get('patient_name'):
        sections.append(f"Patient: {care_plan_data['patient_name']}")
    if care_plan_data.get('age'):
        sections.append(f"Age: {care_plan_data['age']}")
    if care_plan_data.get('sex'):
        sections.append(f"Sex: {care_plan_data['sex']}")
    if care_plan_data.get('diagnosis'):
        sections.append(f"Diagnosis: {care_plan_data['diagnosis']}")
    if care_plan_data.get('procedure'):
        sections.append(f"Procedure: {care_plan_data['procedure']}")
    
    # Medical History
    if care_plan_data.get('pmh'):
        sections.append(f"Past Medical History: {care_plan_data['pmh']}")
    if care_plan_data.get('psh'):
        sections.append(f"Past Surgical History: {care_plan_data['psh']}")
    if care_plan_data.get('meds'):
        sections.append(f"Current Medications: {care_plan_data['meds']}")
    
    # Physical Assessment
    if care_plan_data.get('neuro'):
        sections.append(f"Neurological: {care_plan_data['neuro']}")
    if care_plan_data.get('cardio'):
        sections.append(f"Cardiovascular: {care_plan_data['cardio']}")
    if care_plan_data.get('resp'):
        sections.append(f"Respiratory: {care_plan_data['resp']}")
    
    # Airway Assessment
    if care_plan_data.get('mallampati'):
        sections.append(f"Mallampati: {care_plan_data['mallampati']}")
    if care_plan_data.get('ulbt'):
        sections.append(f"ULBT Grade: {care_plan_data['ulbt']}")
    
    # Laboratory Values
    lab_values = []
    for lab in ['na', 'k', 'cl', 'co2', 'bun', 'cr', 'glu', 'wbc', 'hgb', 'hct', 'plt', 'pt', 'ptt', 'inr']:
        if care_plan_data.get(lab):
            lab_values.append(f"{lab.upper()}: {care_plan_data[lab]}")
    if lab_values:
        sections.append(f"Laboratory Values: {', '.join(lab_values)}")
    
    return "\n".join(sections)

async def index_care_plan_for_rag(care_plan: CarePlan, db: Session):
    """Index care plan content for RAG retrieval"""
    try:
        if not care_plan.exported_text:
            return
        
        # Create vector index entry for the care plan
        embedding = generate_embeddings(care_plan.exported_text)
        
        vector_entry = VectorIndexEntry(
            user_id=care_plan.user_id,
            content_hash=f"care_plan_{care_plan.id}",
            embedding=embedding,
            content=care_plan.exported_text,
            token_count=len(care_plan.exported_text.split()),
            chunk_index=0,
            source_type="care_plan",
            source_id=care_plan.id,
            embedding_model="text-embedding-3-small",
            vector_metadata={
                "care_plan_id": care_plan.id,
                "patient_name": care_plan.patient_name,
                "procedure": care_plan.procedure,
                "diagnosis": care_plan.diagnosis,
                "created_at": care_plan.created_at.isoformat() if care_plan.created_at else None
            }
        )
        
        db.add(vector_entry)
        db.commit()
        
    except Exception as e:
        print(f"Error indexing care plan for RAG: {e}")

def build_care_plan_context(care_plan: CarePlan, db: Session) -> str:
    """Build context for AI generation by searching relevant materials"""
    try:
        # Search for relevant materials using care plan content
        query_text = f"{care_plan.diagnosis} {care_plan.procedure} anesthesia care plan"
        query_embedding = generate_embeddings(query_text)
        
        # Get relevant vector entries (user's materials AND system materials)
        vector_entries = db.query(VectorIndexEntry).filter(
            or_(
                VectorIndexEntry.user_id == care_plan.user_id,
                VectorIndexEntry.user_id == SYSTEM_USER_ID
            )
        ).all()
        
        # Calculate similarity and get top results
        results = []
        for entry in vector_entries:
            similarity = sum(a * b for a, b in zip(query_embedding, entry.embedding))
            results.append({
                "content": entry.content,
                "similarity": similarity,
                "metadata": entry.vector_metadata,
                "is_system": entry.user_id == SYSTEM_USER_ID
            })
        
        # Sort by similarity and get top 5 most relevant chunks
        results.sort(key=lambda x: x["similarity"], reverse=True)
        top_results = results[:5]
        
        context_parts = [f"Patient Care Plan:\n{care_plan.exported_text}"]
        
        if top_results:
            context_parts.append("\nRelevant Medical Literature:")
            for i, result in enumerate(top_results, 1):
                source_name = result["metadata"].get("file_name", "Unknown Source")
                source_label = "System Textbook" if result.get("is_system") else "User Material"
                context_parts.append(f"\n{i}. From {source_name} ({source_label}):\n{result['content'][:300]}...")
        
        return "\n".join(context_parts)
        
    except Exception as e:
        print(f"Error building care plan context: {e}")
        return care_plan.exported_text or ""

@app.post("/chat-test", response_model=ChatOut)
def chat_test(payload: ChatIn, db: Session = Depends(get_db)):
    """Test chat endpoint without authentication"""
    if not client:
        raise HTTPException(status_code=503, detail="OpenAI client not configured")
    
    # Create or reuse a thread
    thread_id = payload.thread_id or str(uuid4())
    
    # Simple conversation without session management for now
    messages = [
        {"role": "system", "content": "You are a helpful AI assistant. You can answer questions about medical topics, general knowledge, current events, and provide practical information. Be helpful and informative in your responses."}
    ]

    messages.append({"role": "user", "content": payload.message})

    try:
        # Use the correct OpenAI model
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=500
        )
        
        reply = response.choices[0].message.content
        
        # For now, just return the response without storing in database
        # to avoid foreign key constraint issues
        
        return ChatOut(reply=reply, thread_id=thread_id)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Chat error: {str(e)}")

# Care Plan API Endpoints
@app.post("/api/care-plans")
async def create_care_plan(
    care_plan_data: dict,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new care plan"""
    try:
        # Create care plan record
        care_plan = CarePlan(
            user_id=current_user['user_id'],
            title=care_plan_data.get('title', 'Untitled Care Plan'),
            patient_name=care_plan_data.get('patient_name'),
            procedure=care_plan_data.get('procedure'),
            diagnosis=care_plan_data.get('diagnosis'),
            
            # Demographics
            age=care_plan_data.get('age'),
            sex=care_plan_data.get('sex'),
            height=care_plan_data.get('height'),
            weight=care_plan_data.get('weight'),
            
            # Vital Signs
            temperature_f=care_plan_data.get('tempF'),
            blood_pressure=care_plan_data.get('bp'),
            heart_rate=care_plan_data.get('hr'),
            respiration_rate=care_plan_data.get('rr'),
            oxygen_saturation=care_plan_data.get('oxygen_saturation'),
            lmp_date=care_plan_data.get('lmp_date'),
            
            # Medical History
            past_medical_history=care_plan_data.get('pmh'),
            past_surgical_history=care_plan_data.get('psh'),
            anesthesia_history=care_plan_data.get('anesthesiaHx'),
            current_medications=care_plan_data.get('meds'),
            alcohol_use=care_plan_data.get('alcohol'),
            substance_use=care_plan_data.get('substance'),
            allergies=care_plan_data.get('allergies'),
            
            # Physical Assessment
            neurological_findings=care_plan_data.get('neuro'),
            heent_findings=care_plan_data.get('heent'),
            respiratory_findings=care_plan_data.get('resp'),
            cardiovascular_findings=care_plan_data.get('cardio'),
            gastrointestinal_findings=care_plan_data.get('gi'),
            genitourinary_findings=care_plan_data.get('gu'),
            endocrine_findings=care_plan_data.get('endo'),
            other_findings=care_plan_data.get('otherFindings'),
            
            # Airway Assessment
            mallampati_class=care_plan_data.get('mallampati'),
            ulbt_grade=care_plan_data.get('ulbt'),
            thyromental_distance=care_plan_data.get('thyromental'),
            interincisor_distance=care_plan_data.get('interincisor'),
            dentition=care_plan_data.get('dentition'),
            neck_assessment=care_plan_data.get('neck'),
            oral_mucosa=care_plan_data.get('oralMucosa'),
            
            # Laboratory Values
            sodium=care_plan_data.get('na'),
            potassium=care_plan_data.get('k'),
            chloride=care_plan_data.get('cl'),
            co2=care_plan_data.get('co2'),
            bun=care_plan_data.get('bun'),
            creatinine=care_plan_data.get('cr'),
            glucose=care_plan_data.get('glu'),
            wbc=care_plan_data.get('wbc'),
            hemoglobin=care_plan_data.get('hgb'),
            hematocrit=care_plan_data.get('hct'),
            platelets=care_plan_data.get('plt'),
            pt=care_plan_data.get('pt'),
            ptt=care_plan_data.get('ptt'),
            inr=care_plan_data.get('inr'),
            abg=care_plan_data.get('abg'),
            other_labs=care_plan_data.get('otherLabs'),
            
            # Imaging/Diagnostic Tests
            ekg=care_plan_data.get('ekg'),
            chest_xray=care_plan_data.get('cxr'),
            echocardiogram=care_plan_data.get('echo'),
            other_imaging=care_plan_data.get('otherImaging'),
            
            # Cultural/Religious Considerations
            cultural_religious_attributes=care_plan_data.get('cultural_religious_attributes'),
            
            # Generate export text for RAG indexing
            exported_text=build_care_plan_text(care_plan_data),
            export_hash=hash(str(care_plan_data))
        )
        
        db.add(care_plan)
        db.commit()
        db.refresh(care_plan)
        
        # Index the care plan content for RAG
        await index_care_plan_for_rag(care_plan, db)
        
        return {
            "success": True,
            "care_plan_id": care_plan.id,
            "message": "Care plan created successfully"
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error creating care plan: {str(e)}")

@app.get("/api/care-plans")
async def get_care_plans(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all care plans for the current user"""
    try:
        care_plans = db.query(CarePlan).filter(
            CarePlan.user_id == current_user['user_id']
        ).order_by(CarePlan.created_at.desc()).all()
        
        return {
            "success": True,
            "care_plans": [
                {
                    "id": cp.id,
                    "title": cp.title,
                    "patient_name": cp.patient_name,
                    "procedure": cp.procedure,
                    "diagnosis": cp.diagnosis,
                    "status": cp.status,
                    "created_at": cp.created_at.isoformat() if cp.created_at else None,
                    "updated_at": cp.updated_at.isoformat() if cp.updated_at else None
                }
                for cp in care_plans
            ]
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching care plans: {str(e)}")

@app.get("/api/care-plans/{care_plan_id}")
async def get_care_plan(
    care_plan_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a specific care plan by ID"""
    try:
        care_plan = db.query(CarePlan).filter(
            CarePlan.id == care_plan_id,
            CarePlan.user_id == current_user['user_id']
        ).first()
        
        if not care_plan:
            raise HTTPException(status_code=404, detail="Care plan not found")
        
        # Update last accessed
        care_plan.last_accessed = func.now()
        db.commit()
        
        return {
            "success": True,
            "care_plan": {
                "id": care_plan.id,
                "title": care_plan.title,
                "patient_name": care_plan.patient_name,
                "procedure": care_plan.procedure,
                "diagnosis": care_plan.diagnosis,
                
                # Demographics
                "age": care_plan.age,
                "sex": care_plan.sex,
                "height": care_plan.height,
                "weight": care_plan.weight,
                
                # Vital Signs
                "tempF": care_plan.temperature_f,
                "bp": care_plan.blood_pressure,
                "hr": care_plan.heart_rate,
                "rr": care_plan.respiration_rate,
                "oxygen_saturation": care_plan.oxygen_saturation,
                "lmp_date": care_plan.lmp_date,
                
                # Medical History
                "pmh": care_plan.past_medical_history,
                "psh": care_plan.past_surgical_history,
                "anesthesiaHx": care_plan.anesthesia_history,
                "meds": care_plan.current_medications,
                "alcohol": care_plan.alcohol_use,
                "substance": care_plan.substance_use,
                "allergies": care_plan.allergies,
                
                # Physical Assessment
                "neuro": care_plan.neurological_findings,
                "heent": care_plan.heent_findings,
                "resp": care_plan.respiratory_findings,
                "cardio": care_plan.cardiovascular_findings,
                "gi": care_plan.gastrointestinal_findings,
                "gu": care_plan.genitourinary_findings,
                "endo": care_plan.endocrine_findings,
                "otherFindings": care_plan.other_findings,
                
                # Airway Assessment
                "mallampati": care_plan.mallampati_class,
                "ulbt": care_plan.ulbt_grade,
                "thyromental": care_plan.thyromental_distance,
                "interincisor": care_plan.interincisor_distance,
                "dentition": care_plan.dentition,
                "neck": care_plan.neck_assessment,
                "oralMucosa": care_plan.oral_mucosa,
                
                # Laboratory Values
                "na": care_plan.sodium,
                "k": care_plan.potassium,
                "cl": care_plan.chloride,
                "co2": care_plan.co2,
                "bun": care_plan.bun,
                "cr": care_plan.creatinine,
                "glu": care_plan.glucose,
                "wbc": care_plan.wbc,
                "hgb": care_plan.hemoglobin,
                "hct": care_plan.hematocrit,
                "plt": care_plan.platelets,
                "pt": care_plan.pt,
                "ptt": care_plan.ptt,
                "inr": care_plan.inr,
                "abg": care_plan.abg,
                "otherLabs": care_plan.other_labs,
                
                # Imaging/Diagnostic Tests
                "ekg": care_plan.ekg,
                "cxr": care_plan.chest_xray,
                "echo": care_plan.echocardiogram,
                "otherImaging": care_plan.other_imaging,
                
                # Cultural/Religious Considerations
                "cultural_religious_attributes": care_plan.cultural_religious_attributes,
                
                # AI-Generated Content
                "ai_recommendations": care_plan.ai_recommendations,
                "risk_assessment": care_plan.risk_assessment,
                "monitoring_plan": care_plan.monitoring_plan,
                "medication_plan": care_plan.medication_plan,
                
                # Metadata
                "status": care_plan.status,
                "version": care_plan.version,
                "created_at": care_plan.created_at.isoformat() if care_plan.created_at else None,
                "updated_at": care_plan.updated_at.isoformat() if care_plan.updated_at else None
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching care plan: {str(e)}")

@app.post("/api/care-plans/{care_plan_id}/generate-ai")
async def generate_ai_recommendations(
    care_plan_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Generate AI recommendations for a care plan using RAG"""
    try:
        care_plan = db.query(CarePlan).filter(
            CarePlan.id == care_plan_id,
            CarePlan.user_id == current_user['user_id']
        ).first()
        
        if not care_plan:
            raise HTTPException(status_code=404, detail="Care plan not found")
        
        if not client:
            raise HTTPException(status_code=503, detail="OpenAI client not configured")
        
        # Build context for AI generation
        context = build_care_plan_context(care_plan, db)
        
        # Generate AI recommendations
        recommendations = await generate_care_plan_recommendations(context, client)
        
        # Update care plan with AI recommendations
        care_plan.ai_recommendations = recommendations.get('anesthesia_plan', '')
        care_plan.risk_assessment = recommendations.get('risk_assessment', '')
        care_plan.monitoring_plan = recommendations.get('monitoring_plan', '')
        care_plan.medication_plan = recommendations.get('medication_plan', '')
        care_plan.rag_context = context
        care_plan.rag_sources = recommendations.get('sources', [])
        care_plan.rag_confidence_score = recommendations.get('confidence_score', 0.0)
        care_plan.updated_at = func.now()
        
        db.commit()
        
        return {
            "success": True,
            "recommendations": recommendations,
            "message": "AI recommendations generated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating AI recommendations: {str(e)}")

@app.delete("/api/care-plans/{care_plan_id}")
async def delete_care_plan(
    care_plan_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a care plan"""
    try:
        care_plan = db.query(CarePlan).filter(
            CarePlan.id == care_plan_id,
            CarePlan.user_id == current_user['user_id']
        ).first()
        
        if not care_plan:
            raise HTTPException(status_code=404, detail="Care plan not found")
        
        # Delete associated RAG entries
        db.query(VectorIndexEntry).filter(
            VectorIndexEntry.source_id == care_plan_id,
            VectorIndexEntry.source_type == "care_plan"
        ).delete()
        
        # Delete care plan
        db.delete(care_plan)
        db.commit()
        
        return {
            "success": True,
            "message": "Care plan deleted successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error deleting care plan: {str(e)}")

# Helper functions for care plan operations
def build_care_plan_text(care_plan_data: dict) -> str:
    """Build text representation of care plan for RAG indexing"""
    sections = []
    
    # Patient Information
    if care_plan_data.get('patient_name'):
        sections.append(f"Patient: {care_plan_data['patient_name']}")
    if care_plan_data.get('age'):
        sections.append(f"Age: {care_plan_data['age']}")
    if care_plan_data.get('sex'):
        sections.append(f"Sex: {care_plan_data['sex']}")
    if care_plan_data.get('diagnosis'):
        sections.append(f"Diagnosis: {care_plan_data['diagnosis']}")
    if care_plan_data.get('procedure'):
        sections.append(f"Procedure: {care_plan_data['procedure']}")
    
    # Medical History
    if care_plan_data.get('pmh'):
        sections.append(f"Past Medical History: {care_plan_data['pmh']}")
    if care_plan_data.get('psh'):
        sections.append(f"Past Surgical History: {care_plan_data['psh']}")
    if care_plan_data.get('meds'):
        sections.append(f"Current Medications: {care_plan_data['meds']}")
    
    # Physical Assessment
    if care_plan_data.get('neuro'):
        sections.append(f"Neurological: {care_plan_data['neuro']}")
    if care_plan_data.get('cardio'):
        sections.append(f"Cardiovascular: {care_plan_data['cardio']}")
    if care_plan_data.get('resp'):
        sections.append(f"Respiratory: {care_plan_data['resp']}")
    
    # Airway Assessment
    if care_plan_data.get('mallampati'):
        sections.append(f"Mallampati: {care_plan_data['mallampati']}")
    if care_plan_data.get('ulbt'):
        sections.append(f"ULBT Grade: {care_plan_data['ulbt']}")
    
    # Laboratory Values
    lab_values = []
    for lab in ['na', 'k', 'cl', 'co2', 'bun', 'cr', 'glu', 'wbc', 'hgb', 'hct', 'plt', 'pt', 'ptt', 'inr']:
        if care_plan_data.get(lab):
            lab_values.append(f"{lab.upper()}: {care_plan_data[lab]}")
    if lab_values:
        sections.append(f"Laboratory Values: {', '.join(lab_values)}")
    
    return "\n".join(sections)

async def index_care_plan_for_rag(care_plan: CarePlan, db: Session):
    """Index care plan content for RAG retrieval"""
    try:
        if not care_plan.exported_text:
            return
        
        # Create vector index entry for the care plan
        embedding = generate_embeddings(care_plan.exported_text)
        
        vector_entry = VectorIndexEntry(
            user_id=care_plan.user_id,
            content_hash=f"care_plan_{care_plan.id}",
            embedding=embedding,
            content=care_plan.exported_text,
            token_count=len(care_plan.exported_text.split()),
            chunk_index=0,
            source_type="care_plan",
            source_id=care_plan.id,
            embedding_model="text-embedding-3-small",
            vector_metadata={
                "care_plan_id": care_plan.id,
                "patient_name": care_plan.patient_name,
                "procedure": care_plan.procedure,
                "diagnosis": care_plan.diagnosis,
                "created_at": care_plan.created_at.isoformat() if care_plan.created_at else None
            }
        )
        
        db.add(vector_entry)
        db.commit()
        
    except Exception as e:
        print(f"Error indexing care plan for RAG: {e}")

def build_care_plan_context(care_plan: CarePlan, db: Session) -> str:
    """Build context for AI generation by searching relevant materials"""
    try:
        # Search for relevant materials using care plan content
        query_text = f"{care_plan.diagnosis} {care_plan.procedure} anesthesia care plan"
        query_embedding = generate_embeddings(query_text)
        
        # Get relevant vector entries (user's materials AND system materials)
        vector_entries = db.query(VectorIndexEntry).filter(
            or_(
                VectorIndexEntry.user_id == care_plan.user_id,
                VectorIndexEntry.user_id == SYSTEM_USER_ID
            )
        ).all()
        
        # Calculate similarity and get top results
        results = []
        for entry in vector_entries:
            similarity = sum(a * b for a, b in zip(query_embedding, entry.embedding))
            results.append({
                "content": entry.content,
                "similarity": similarity,
                "metadata": entry.vector_metadata,
                "is_system": entry.user_id == SYSTEM_USER_ID
            })
        
        # Sort by similarity and get top 5 most relevant chunks
        results.sort(key=lambda x: x["similarity"], reverse=True)
        top_results = results[:5]
        
        context_parts = [f"Patient Care Plan:\n{care_plan.exported_text}"]
        
        if top_results:
            context_parts.append("\nRelevant Medical Literature:")
            for i, result in enumerate(top_results, 1):
                source_name = result["metadata"].get("file_name", "Unknown Source")
                source_label = "System Textbook" if result.get("is_system") else "User Material"
                context_parts.append(f"\n{i}. From {source_name} ({source_label}):\n{result['content'][:300]}...")
        
        return "\n".join(context_parts)
        
    except Exception as e:
        print(f"Error building care plan context: {e}")
        return care_plan.exported_text or ""

async def generate_care_plan_recommendations(context: str, client) -> dict:
    """Generate AI recommendations for care plan"""
    try:
        system_prompt = """You are an expert anesthesiologist AI assistant. Based on the patient information and relevant medical literature provided, generate comprehensive anesthesia care plan recommendations.

Please provide:
1. Anesthesia Plan: Detailed anesthesia approach including induction, maintenance, and emergence
2. Risk Assessment: Analysis of patient-specific risks and complications
3. Monitoring Plan: Required monitoring during the procedure
4. Medication Plan: Specific medications and dosages

Be specific, evidence-based, and consider the patient's comorbidities and procedure requirements."""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": context}
        ]
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=1500
        )
        
        ai_response = response.choices[0].message.content
        
        # Parse the response into structured recommendations
        recommendations = {
            "anesthesia_plan": ai_response,
            "risk_assessment": "",
            "monitoring_plan": "",
            "medication_plan": "",
            "sources": [],
            "confidence_score": 0.8
        }
        
        return recommendations
        
    except Exception as e:
        print(f"Error generating AI recommendations: {e}")
        return {
            "anesthesia_plan": "Error generating recommendations",
            "risk_assessment": "",
            "monitoring_plan": "",
            "medication_plan": "",
            "sources": [],
            "confidence_score": 0.0
        }

# File Upload Endpoints
@app.post("/api/upload")
async def upload_file(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload and process a file (PDF, DOCX, TXT) for RAG"""
    
    # Validate file type
    allowed_types = ["pdf", "docx", "doc", "txt"]
    file_extension = file.filename.split('.')[-1].lower() if '.' in file.filename else ""
    
    if file_extension not in allowed_types:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_types)}"
        )
    
    # Read file content
    try:
        file_content = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")
    
    # Generate unique file path
    file_id = str(uuid4())
    s3_key = f"uploads/{current_user['user_id']}/{file_id}_{file.filename}"
    
    # Store file path (S3 if configured, otherwise None for text-only storage)
    file_path = None
    if s3_client:
        try:
            s3_client.put_object(
                Bucket=S3_BUCKET_NAME,
                Key=s3_key,
                Body=file_content,
                ContentType=file.content_type or "application/octet-stream"
            )
            file_path = f"s3://{S3_BUCKET_NAME}/{s3_key}"
            print(f"File uploaded to S3: {file_path}")
        except Exception as e:
            print(f"S3 upload failed: {e}")
            file_path = None
    else:
        print("S3 not configured - storing text content only (no file download available)")
        file_path = None
    
    # Extract text from file
    try:
        extracted_text = extract_text_from_file(file_content, file_extension)
        print(f"Extracted text length: {len(extracted_text)} characters")
        print(f"First 200 chars: {extracted_text[:200]}")
    except Exception as e:
        print(f"Text extraction error: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")
    
    # Create material record in database
    material = Material(
        user_id=current_user['user_id'],
        title=file.filename,
        file_type=file_extension,
        file_path=file_path,
        file_size=len(file_content),
        status="processing",
        processing_progress=0
    )
    
    db.add(material)
    db.commit()
    db.refresh(material)
    
    # Process text for RAG (chunk and create embeddings)
    try:
        chunks = chunk_text(extracted_text)
        
        # Create vector index entries for each chunk
        for i, chunk in enumerate(chunks):
            try:
                embedding = generate_embeddings(chunk)
                
                vector_entry = VectorIndexEntry(
                    user_id=current_user['user_id'],
                    content_hash=f"{file_id}_{i}",
                    embedding=embedding,
                    content=chunk,
                    token_count=len(chunk.split()),  # Approximate token count
                    chunk_index=i,
                    source_type="material",
                    source_id=material.id,
                    embedding_model="text-embedding-3-small",
                    vector_metadata={
                        "file_name": file.filename,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "file_type": file_extension,
                        "file_size": len(file_content)
                    }
                )
                
                db.add(vector_entry)
                
            except Exception as e:
                error_msg = str(e)
                print(f"Error creating embedding for chunk {i}: {error_msg}")
                # Log specific error types for debugging
                if "403" in error_msg:
                    print(f"⚠️  OpenAI API 403 error for chunk {i}. This chunk will be skipped. Check your API key permissions/quota.")
                elif "401" in error_msg:
                    print(f"⚠️  OpenAI API 401 error for chunk {i}. Invalid API key. Please check your OPENAI_API_KEY.")
                continue
        
        db.commit()
        
        # Update material status
        material.status = "processed"
        material.processing_progress = 100
        material.chunk_count = len(chunks)
        material.total_tokens = sum(len(chunk.split()) for chunk in chunks)
        material.extracted_text = extracted_text
        material.processed_at = func.now()
        db.commit()
        
        # Cache the extracted text
        cache_text(material.id, extracted_text)
        
        return {
            "success": True,
            "material_id": material.id,
            "file_name": file.filename,
            "file_type": file_extension,
            "chunks_created": len(chunks),
            "text_length": len(extracted_text),
            "file_path": file_path,
            "message": f"File processed successfully. Created {len(chunks)} chunks for RAG."
        }
        
    except Exception as e:
        # Update material status to failed
        material.status = "failed"
        material.processing_error = str(e)
        db.commit()
        raise HTTPException(status_code=500, detail=f"Error processing file for RAG: {str(e)}")

@app.post("/api/admin/upload-system-material")
async def upload_system_material(
    file: UploadFile = File(...),
    x_admin_key: Optional[str] = Header(None, alias="X-Admin-Key"),
    db: Session = Depends(get_db)
):
    """Admin endpoint to upload system materials accessible to all users"""
    
    # Optional: Check admin key (set ADMIN_API_KEY in environment)
    admin_api_key = os.getenv("ADMIN_API_KEY")
    if admin_api_key:
        if not x_admin_key or x_admin_key != admin_api_key:
            raise HTTPException(status_code=403, detail="Invalid admin key")
    
    # Validate file type
    allowed_types = ["pdf", "docx", "doc", "txt"]
    file_extension = file.filename.split('.')[-1].lower() if '.' in file.filename else ""
    
    if file_extension not in allowed_types:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_types)}"
        )
    
    # Read file content
    try:
        file_content = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")
    
    # Generate unique file path
    file_id = str(uuid4())
    s3_key = f"system-materials/{file_id}_{file.filename}"
    
    # Store file path (S3 if configured, otherwise None for text-only storage)
    file_path = None
    if s3_client:
        try:
            s3_client.put_object(
                Bucket=S3_BUCKET_NAME,
                Key=s3_key,
                Body=file_content,
                ContentType=file.content_type or "application/octet-stream"
            )
            file_path = f"s3://{S3_BUCKET_NAME}/{s3_key}"
            print(f"System material uploaded to S3: {file_path}")
        except Exception as e:
            print(f"S3 upload failed: {e}")
            file_path = None
    else:
        print("S3 not configured - storing text content only (no file download available)")
        file_path = None
    
    # Extract text from file
    try:
        extracted_text = extract_text_from_file(file_content, file_extension)
        print(f"Extracted text length: {len(extracted_text)} characters")
    except Exception as e:
        print(f"Text extraction error: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")
    
    # Check if material with same title already exists (prevent duplicates)
    existing_material = db.query(Material).filter(
        Material.user_id == SYSTEM_USER_ID,
        Material.title == file.filename,
        Material.status == "processed"
    ).first()
    
    if existing_material:
        raise HTTPException(
            status_code=400, 
            detail=f"System material with title '{file.filename}' already exists"
        )
    
    # Create material record in database with SYSTEM_USER_ID
    material = Material(
        user_id=SYSTEM_USER_ID,  # System materials accessible to all users
        title=file.filename,
        file_type=file_extension,
        file_path=file_path,
        file_size=len(file_content),
        status="processing",
        processing_progress=0
    )
    
    db.add(material)
    db.commit()
    db.refresh(material)
    
    # Process text for RAG (chunk and create embeddings) - run in background
    # Return immediately to avoid timeout for large files
    import threading
    
    def process_material_background():
        """Process material in background to avoid timeout"""
        # Get a new database session for background processing
        from database import get_session_local
        background_db = get_session_local()()
        
        try:
            # Reload material in background session
            bg_material = background_db.query(Material).filter(Material.id == material.id).first()
            if not bg_material:
                print(f"Material {material.id} not found in background processing")
                return
            
            chunks = chunk_text(extracted_text)
            total_chunks = len(chunks)
            
            # Create vector index entries for each chunk
            successful_chunks = 0
            for i, chunk in enumerate(chunks):
                try:
                    # Update progress
                    bg_material.processing_progress = int((i / total_chunks) * 100)
                    background_db.commit()
                    
                    embedding = generate_embeddings(chunk)
                    
                    vector_entry = VectorIndexEntry(
                        user_id=SYSTEM_USER_ID,  # System materials accessible to all users
                        content_hash=f"{file_id}_{i}",
                        embedding=embedding,
                        content=chunk,
                        token_count=len(chunk.split()),  # Approximate token count
                        chunk_index=i,
                        source_type="material",
                        source_id=bg_material.id,
                        embedding_model="text-embedding-3-small",
                        vector_metadata={
                            "file_name": file.filename,
                            "chunk_index": i,
                            "total_chunks": total_chunks,
                            "file_type": file_extension,
                            "file_size": len(file_content),
                            "is_system": True
                        }
                    )
                    
                    background_db.add(vector_entry)
                    successful_chunks += 1
                    
                    # Commit every 10 chunks to avoid long transactions
                    if (i + 1) % 10 == 0:
                        background_db.commit()
                    
                except Exception as e:
                    error_msg = str(e)
                    print(f"Error creating embedding for chunk {i}: {error_msg}")
                    # Log specific error types for debugging
                    if "403" in error_msg:
                        print(f"⚠️  OpenAI API 403 error for chunk {i}. This chunk will be skipped. Check your API key permissions/quota.")
                    elif "401" in error_msg:
                        print(f"⚠️  OpenAI API 401 error for chunk {i}. Invalid API key. Please check your OPENAI_API_KEY.")
                    continue
            
            background_db.commit()
            
            # Update material status
            bg_material.status = "processed"
            bg_material.processing_progress = 100
            bg_material.chunk_count = successful_chunks
            bg_material.total_tokens = sum(len(chunk.split()) for chunk in chunks)
            bg_material.extracted_text = extracted_text
            bg_material.processed_at = func.now()
            background_db.commit()
            
            # Cache the extracted text
            cache_text(bg_material.id, extracted_text)
            
            print(f"✓ Successfully processed system material: {file.filename} ({successful_chunks} chunks)")
            
        except Exception as e:
            # Update material status to failed
            try:
                bg_material = background_db.query(Material).filter(Material.id == material.id).first()
                if bg_material:
                    bg_material.status = "failed"
                    bg_material.processing_error = str(e)
                    background_db.commit()
            except:
                pass
            print(f"✗ Error processing system material {file.filename}: {str(e)}")
        finally:
            background_db.close()
    
    # Start background processing in a separate thread (fire and forget)
    thread = threading.Thread(target=process_material_background, daemon=True)
    thread.start()
    
    # Return immediately with processing status
    return {
        "success": True,
        "material_id": material.id,
        "file_name": file.filename,
        "file_type": file_extension,
        "file_size": len(file_content),
        "file_path": file_path,
        "status": "processing",
        "message": f"System material uploaded successfully. Processing in background. Check status later."
    }

@app.get("/api/materials")
def get_user_materials(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all materials uploaded by the current user"""
    
    materials = db.query(Material).filter(
        Material.user_id == current_user['user_id']
    ).order_by(Material.uploaded_at.desc()).all()
    
    result_materials = []
    for material in materials:
        # Try to get from cache first
        cached_text = get_cached_text(material.id)
        
        # If not in cache, use database value and cache it
        extracted_text = cached_text if cached_text is not None else material.extracted_text
        
        # Cache the text if it wasn't cached and exists
        if cached_text is None and extracted_text:
            cache_text(material.id, extracted_text)
        
        result_materials.append({
            "id": material.id,
            "title": material.title,
            "file_type": material.file_type,
            "file_path": material.file_path,
            "file_size": material.file_size,
            "status": material.status,
            "processing_progress": material.processing_progress,
            "processing_error": material.processing_error,
            "extracted_text": extracted_text,
            "chunk_count": material.chunk_count,
            "total_tokens": material.total_tokens,
            "embedding_model": material.embedding_model,
            "uploaded_at": material.uploaded_at.isoformat(),
            "processed_at": material.processed_at.isoformat() if material.processed_at else None,
            "last_accessed": material.last_accessed.isoformat() if material.last_accessed else None
        })
    
    return {
        "materials": result_materials
    }

@app.delete("/api/materials/{material_id}")
def delete_material(
    material_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a material and its associated vector entries"""
    
    # Find the material
    material = db.query(Material).filter(
        Material.id == material_id,
        Material.user_id == current_user['user_id']
    ).first()
    
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")
    
    # Invalidate caches before deleting
    invalidate_cache(material_id)
    invalidate_vector_cache(material_id)
    
    # Delete associated vector entries
    vector_entries = db.query(VectorIndexEntry).filter(
        VectorIndexEntry.source_id == material_id,
        VectorIndexEntry.source_type == "material"
    ).all()
    
    for entry in vector_entries:
        db.delete(entry)
    
    # Delete the material
    db.delete(material)
    db.commit()
    
    return {"success": True, "message": "Material deleted successfully"}

@app.get("/api/search")
def search_materials(
    query: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
    limit: int = 5
):
    """Search through user's materials and system materials using vector similarity"""
    
    if not client:
        raise HTTPException(status_code=503, detail="OpenAI client not configured")
    
    try:
        # Generate embedding for the query
        query_embedding = generate_embeddings(query)
        
        # Get user's materials AND system materials (accessible to all users)
        user_materials = db.query(Material).filter(
            or_(
                Material.user_id == current_user['user_id'],
                Material.user_id == SYSTEM_USER_ID
            ),
            Material.status == "processed"
        ).all()
        
        material_ids = [m.id for m in user_materials]
        
        if not material_ids:
            return {"results": [], "message": "No processed materials found"}
        
        # Get vector entries for user's materials and system materials
        vector_entries = db.query(VectorIndexEntry).filter(
            VectorIndexEntry.source_id.in_(material_ids),
            VectorIndexEntry.source_type == "material",
            or_(
                VectorIndexEntry.user_id == current_user['user_id'],
                VectorIndexEntry.user_id == SYSTEM_USER_ID
            )
        ).all()
        
        # Calculate similarity scores (simplified cosine similarity)
        results = []
        for entry in vector_entries:
            # Simple dot product for similarity (in production, use proper cosine similarity)
            similarity = sum(a * b for a, b in zip(query_embedding, entry.embedding))
            results.append({
                "content": entry.content,
                "similarity": similarity,
                "metadata": entry.vector_metadata,
                "source_id": entry.source_id,
                "is_system": entry.user_id == SYSTEM_USER_ID
            })
        
        # Sort by similarity and return top results
        results.sort(key=lambda x: x["similarity"], reverse=True)
        
        return {
            "results": results[:limit],
            "query": query,
            "total_found": len(results)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Search error: {str(e)}")

@app.post("/chat-rag", response_model=ChatOut)
def chat_with_rag(payload: ChatIn, current_user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Enhanced chat endpoint with RAG integration - searches user's materials for relevant context"""
    if not client:
        raise HTTPException(status_code=503, detail="OpenAI client not configured")
    
    # Create or reuse a thread
    thread_id = payload.thread_id or str(uuid4())
    
    # RAG: Search user's materials AND system materials for relevant context
    relevant_context = ""
    try:
        # Generate embedding for the user's question
        query_embedding = generate_embeddings(payload.message)
        
        # Get user's processed materials AND system materials (accessible to all users)
        user_materials = db.query(Material).filter(
            or_(
                Material.user_id == current_user['user_id'],
                Material.user_id == SYSTEM_USER_ID
            ),
            Material.status == "processed"
        ).all()
        
        if user_materials:
            material_ids = [m.id for m in user_materials]
            
            # Get vector entries for user's materials and system materials
            vector_entries = db.query(VectorIndexEntry).filter(
                VectorIndexEntry.source_id.in_(material_ids),
                VectorIndexEntry.source_type == "material",
                or_(
                    VectorIndexEntry.user_id == current_user['user_id'],
                    VectorIndexEntry.user_id == SYSTEM_USER_ID
                )
            ).all()
            
            # Calculate similarity scores and get top results
            results = []
            for entry in vector_entries:
                similarity = sum(a * b for a, b in zip(query_embedding, entry.embedding))
                results.append({
                    "content": entry.content,
                    "similarity": similarity,
                    "metadata": entry.vector_metadata,
                    "is_system": entry.user_id == SYSTEM_USER_ID
                })
            
            # Sort by similarity and get top 5 most relevant chunks (increased from 3 to include system materials)
            results.sort(key=lambda x: x["similarity"], reverse=True)
            top_results = results[:5]
            
            if top_results:
                relevant_context = "\n\nRelevant information from available materials:\n"
                for i, result in enumerate(top_results, 1):
                    file_name = result["metadata"].get("file_name", "Unknown")
                    source_label = "System Textbook" if result.get("is_system") else "Your Upload"
                    relevant_context += f"\n{i}. From {file_name} ({source_label}):\n{result['content'][:500]}...\n"
                
                # Update access tracking (only for user's materials, not system materials)
                for entry in vector_entries:
                    if entry.user_id != SYSTEM_USER_ID and any(entry.content == result["content"] for result in top_results):
                        entry.last_accessed = func.now()
                        entry.access_count += 1
                db.commit()
    
    except Exception as e:
        error_msg = str(e)
        print(f"RAG search error: {error_msg}")
        # Log more specific error information
        if "403" in error_msg or "forbidden" in error_msg.lower():
            print("⚠️  OpenAI API key issue detected (403). Chatbot will work without RAG context. Please check your API key configuration.")
        elif "401" in error_msg or "invalid_api_key" in error_msg.lower():
            print("⚠️  Invalid OpenAI API key detected. Chatbot will work without RAG context. Please update your OPENAI_API_KEY in .env file.")
        # Continue without RAG context if search fails - chatbot will still work
    
    # Build messages with RAG context
    system_prompt = f"""You are a helpful AI assistant for Clyvara, a medical education platform. You can answer questions about medical topics, general knowledge, current events, and provide practical information.

{relevant_context}

When referencing information from uploaded materials, mention the source file name. Be helpful and informative in your responses."""
    
    messages = [
        {"role": "system", "content": system_prompt}
    ]

    messages.append({"role": "user", "content": payload.message})

    try:
        # Use the correct OpenAI model
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=500
        )
        
        reply = response.choices[0].message.content
        
        # Store the chat message in database
        try:
            # First create a user session if it doesn't exist
            existing_session = db.query(UserSession).filter(UserSession.session_id == thread_id).first()
            if not existing_session:
                user_session = UserSession(
                    session_id=thread_id,
                    user_id=current_user['user_id'],
                    is_active=True
                )
                db.add(user_session)
                db.commit()
            
            # Store the chat message
            chat_message = ChatMessage(
                session_id=thread_id,
                message_type="user",
                message_content={"message": payload.message},
                user_id=current_user['user_id'],
                response_time_ms=100,  # Placeholder
                message_length=len(payload.message)
            )
            
            db.add(chat_message)
            db.commit()
            
        except Exception as db_error:
            # Log the error but don't fail the chat
            print(f"Database storage error: {db_error}")
        
        return ChatOut(reply=reply, thread_id=thread_id)
        
    except Exception as e:
        error_msg = str(e)
        # Provide more specific error messages for common OpenAI API issues
        if "401" in error_msg or "invalid_api_key" in error_msg.lower():
            raise HTTPException(
                status_code=500,
                detail="Chat error: Invalid OpenAI API key. Please check your OPENAI_API_KEY in the .env file."
            )
        elif "403" in error_msg or "forbidden" in error_msg.lower():
            raise HTTPException(
                status_code=500,
                detail="Chat error: API key access denied (403). This may be due to insufficient quota, expired key, or restricted permissions. Please check your OpenAI account."
            )
        elif "429" in error_msg or "rate limit" in error_msg.lower():
            raise HTTPException(
                status_code=500,
                detail="Chat error: Rate limit exceeded. Please wait a moment and try again."
            )
        else:
            raise HTTPException(status_code=500, detail=f"Chat error: {error_msg}")

# Care Plan API Endpoints
@app.post("/api/care-plans")
async def create_care_plan(
    care_plan_data: dict,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new care plan"""
    try:
        # Create care plan record
        care_plan = CarePlan(
            user_id=current_user['user_id'],
            title=care_plan_data.get('title', 'Untitled Care Plan'),
            patient_name=care_plan_data.get('patient_name'),
            procedure=care_plan_data.get('procedure'),
            diagnosis=care_plan_data.get('diagnosis'),
            
            # Demographics
            age=care_plan_data.get('age'),
            sex=care_plan_data.get('sex'),
            height=care_plan_data.get('height'),
            weight=care_plan_data.get('weight'),
            
            # Vital Signs
            temperature_f=care_plan_data.get('tempF'),
            blood_pressure=care_plan_data.get('bp'),
            heart_rate=care_plan_data.get('hr'),
            respiration_rate=care_plan_data.get('rr'),
            oxygen_saturation=care_plan_data.get('oxygen_saturation'),
            lmp_date=care_plan_data.get('lmp_date'),
            
            # Medical History
            past_medical_history=care_plan_data.get('pmh'),
            past_surgical_history=care_plan_data.get('psh'),
            anesthesia_history=care_plan_data.get('anesthesiaHx'),
            current_medications=care_plan_data.get('meds'),
            alcohol_use=care_plan_data.get('alcohol'),
            substance_use=care_plan_data.get('substance'),
            allergies=care_plan_data.get('allergies'),
            
            # Physical Assessment
            neurological_findings=care_plan_data.get('neuro'),
            heent_findings=care_plan_data.get('heent'),
            respiratory_findings=care_plan_data.get('resp'),
            cardiovascular_findings=care_plan_data.get('cardio'),
            gastrointestinal_findings=care_plan_data.get('gi'),
            genitourinary_findings=care_plan_data.get('gu'),
            endocrine_findings=care_plan_data.get('endo'),
            other_findings=care_plan_data.get('otherFindings'),
            
            # Airway Assessment
            mallampati_class=care_plan_data.get('mallampati'),
            ulbt_grade=care_plan_data.get('ulbt'),
            thyromental_distance=care_plan_data.get('thyromental'),
            interincisor_distance=care_plan_data.get('interincisor'),
            dentition=care_plan_data.get('dentition'),
            neck_assessment=care_plan_data.get('neck'),
            oral_mucosa=care_plan_data.get('oralMucosa'),
            
            # Laboratory Values
            sodium=care_plan_data.get('na'),
            potassium=care_plan_data.get('k'),
            chloride=care_plan_data.get('cl'),
            co2=care_plan_data.get('co2'),
            bun=care_plan_data.get('bun'),
            creatinine=care_plan_data.get('cr'),
            glucose=care_plan_data.get('glu'),
            wbc=care_plan_data.get('wbc'),
            hemoglobin=care_plan_data.get('hgb'),
            hematocrit=care_plan_data.get('hct'),
            platelets=care_plan_data.get('plt'),
            pt=care_plan_data.get('pt'),
            ptt=care_plan_data.get('ptt'),
            inr=care_plan_data.get('inr'),
            abg=care_plan_data.get('abg'),
            other_labs=care_plan_data.get('otherLabs'),
            
            # Imaging/Diagnostic Tests
            ekg=care_plan_data.get('ekg'),
            chest_xray=care_plan_data.get('cxr'),
            echocardiogram=care_plan_data.get('echo'),
            other_imaging=care_plan_data.get('otherImaging'),
            
            # Cultural/Religious Considerations
            cultural_religious_attributes=care_plan_data.get('cultural_religious_attributes'),
            
            # Generate export text for RAG indexing
            exported_text=build_care_plan_text(care_plan_data),
            export_hash=hash(str(care_plan_data))
        )
        
        db.add(care_plan)
        db.commit()
        db.refresh(care_plan)
        
        # Index the care plan content for RAG
        await index_care_plan_for_rag(care_plan, db)
        
        return {
            "success": True,
            "care_plan_id": care_plan.id,
            "message": "Care plan created successfully"
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error creating care plan: {str(e)}")

@app.get("/api/care-plans")
async def get_care_plans(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all care plans for the current user"""
    try:
        care_plans = db.query(CarePlan).filter(
            CarePlan.user_id == current_user['user_id']
        ).order_by(CarePlan.created_at.desc()).all()
        
        return {
            "success": True,
            "care_plans": [
                {
                    "id": cp.id,
                    "title": cp.title,
                    "patient_name": cp.patient_name,
                    "procedure": cp.procedure,
                    "diagnosis": cp.diagnosis,
                    "status": cp.status,
                    "created_at": cp.created_at.isoformat() if cp.created_at else None,
                    "updated_at": cp.updated_at.isoformat() if cp.updated_at else None
                }
                for cp in care_plans
            ]
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching care plans: {str(e)}")

@app.get("/api/care-plans/{care_plan_id}")
async def get_care_plan(
    care_plan_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a specific care plan by ID"""
    try:
        care_plan = db.query(CarePlan).filter(
            CarePlan.id == care_plan_id,
            CarePlan.user_id == current_user['user_id']
        ).first()
        
        if not care_plan:
            raise HTTPException(status_code=404, detail="Care plan not found")
        
        # Update last accessed
        care_plan.last_accessed = func.now()
        db.commit()
        
        return {
            "success": True,
            "care_plan": {
                "id": care_plan.id,
                "title": care_plan.title,
                "patient_name": care_plan.patient_name,
                "procedure": care_plan.procedure,
                "diagnosis": care_plan.diagnosis,
                
                # Demographics
                "age": care_plan.age,
                "sex": care_plan.sex,
                "height": care_plan.height,
                "weight": care_plan.weight,
                
                # Vital Signs
                "tempF": care_plan.temperature_f,
                "bp": care_plan.blood_pressure,
                "hr": care_plan.heart_rate,
                "rr": care_plan.respiration_rate,
                "oxygen_saturation": care_plan.oxygen_saturation,
                "lmp_date": care_plan.lmp_date,
                
                # Medical History
                "pmh": care_plan.past_medical_history,
                "psh": care_plan.past_surgical_history,
                "anesthesiaHx": care_plan.anesthesia_history,
                "meds": care_plan.current_medications,
                "alcohol": care_plan.alcohol_use,
                "substance": care_plan.substance_use,
                "allergies": care_plan.allergies,
                
                # Physical Assessment
                "neuro": care_plan.neurological_findings,
                "heent": care_plan.heent_findings,
                "resp": care_plan.respiratory_findings,
                "cardio": care_plan.cardiovascular_findings,
                "gi": care_plan.gastrointestinal_findings,
                "gu": care_plan.genitourinary_findings,
                "endo": care_plan.endocrine_findings,
                "otherFindings": care_plan.other_findings,
                
                # Airway Assessment
                "mallampati": care_plan.mallampati_class,
                "ulbt": care_plan.ulbt_grade,
                "thyromental": care_plan.thyromental_distance,
                "interincisor": care_plan.interincisor_distance,
                "dentition": care_plan.dentition,
                "neck": care_plan.neck_assessment,
                "oralMucosa": care_plan.oral_mucosa,
                
                # Laboratory Values
                "na": care_plan.sodium,
                "k": care_plan.potassium,
                "cl": care_plan.chloride,
                "co2": care_plan.co2,
                "bun": care_plan.bun,
                "cr": care_plan.creatinine,
                "glu": care_plan.glucose,
                "wbc": care_plan.wbc,
                "hgb": care_plan.hemoglobin,
                "hct": care_plan.hematocrit,
                "plt": care_plan.platelets,
                "pt": care_plan.pt,
                "ptt": care_plan.ptt,
                "inr": care_plan.inr,
                "abg": care_plan.abg,
                "otherLabs": care_plan.other_labs,
                
                # Imaging/Diagnostic Tests
                "ekg": care_plan.ekg,
                "cxr": care_plan.chest_xray,
                "echo": care_plan.echocardiogram,
                "otherImaging": care_plan.other_imaging,
                
                # Cultural/Religious Considerations
                "cultural_religious_attributes": care_plan.cultural_religious_attributes,
                
                # AI-Generated Content
                "ai_recommendations": care_plan.ai_recommendations,
                "risk_assessment": care_plan.risk_assessment,
                "monitoring_plan": care_plan.monitoring_plan,
                "medication_plan": care_plan.medication_plan,
                
                # Metadata
                "status": care_plan.status,
                "version": care_plan.version,
                "created_at": care_plan.created_at.isoformat() if care_plan.created_at else None,
                "updated_at": care_plan.updated_at.isoformat() if care_plan.updated_at else None
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching care plan: {str(e)}")

@app.post("/api/care-plans/{care_plan_id}/generate-ai")
async def generate_ai_recommendations(
    care_plan_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Generate AI recommendations for a care plan using RAG"""
    try:
        care_plan = db.query(CarePlan).filter(
            CarePlan.id == care_plan_id,
            CarePlan.user_id == current_user['user_id']
        ).first()
        
        if not care_plan:
            raise HTTPException(status_code=404, detail="Care plan not found")
        
        if not client:
            raise HTTPException(status_code=503, detail="OpenAI client not configured")
        
        # Build context for AI generation
        context = build_care_plan_context(care_plan, db)
        
        # Generate AI recommendations
        recommendations = await generate_care_plan_recommendations(context, client)
        
        # Update care plan with AI recommendations
        care_plan.ai_recommendations = recommendations.get('anesthesia_plan', '')
        care_plan.risk_assessment = recommendations.get('risk_assessment', '')
        care_plan.monitoring_plan = recommendations.get('monitoring_plan', '')
        care_plan.medication_plan = recommendations.get('medication_plan', '')
        care_plan.rag_context = context
        care_plan.rag_sources = recommendations.get('sources', [])
        care_plan.rag_confidence_score = recommendations.get('confidence_score', 0.0)
        care_plan.updated_at = func.now()
        
        db.commit()
        
        return {
            "success": True,
            "recommendations": recommendations,
            "message": "AI recommendations generated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating AI recommendations: {str(e)}")

@app.delete("/api/care-plans/{care_plan_id}")
async def delete_care_plan(
    care_plan_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a care plan"""
    try:
        care_plan = db.query(CarePlan).filter(
            CarePlan.id == care_plan_id,
            CarePlan.user_id == current_user['user_id']
        ).first()
        
        if not care_plan:
            raise HTTPException(status_code=404, detail="Care plan not found")
        
        # Delete associated RAG entries
        db.query(VectorIndexEntry).filter(
            VectorIndexEntry.source_id == care_plan_id,
            VectorIndexEntry.source_type == "care_plan"
        ).delete()
        
        # Delete care plan
        db.delete(care_plan)
        db.commit()
        
        return {
            "success": True,
            "message": "Care plan deleted successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error deleting care plan: {str(e)}")

# Learning Plan Question Generation with RAG
class GenerateQuestionsRequest(BaseModel):
    case_study: Optional[str] = None
    topic: Optional[str] = None
    num_questions: int = 3

@app.post("/api/learning-plan/generate-questions")
async def generate_learning_plan_questions(
    request: GenerateQuestionsRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Generate knowledge check questions using RAG from system PDFs and user-uploaded materials"""
    if not client:
        raise HTTPException(status_code=503, detail="OpenAI client not configured")
    
    try:
        # Build query text from case study or topic, or use general query for system materials
        query_text = ""
        use_general_query = False
        
        if request.case_study:
            query_text = f"Case Study: {request.case_study}"
        elif request.topic:
            query_text = f"Topic: {request.topic}"
        else:
            # No case study or topic provided - use general medical education query
            query_text = "medical education anesthesia nursing clinical concepts"
            use_general_query = True
        
        # RAG: Search user's materials AND system materials for relevant context
        relevant_context = ""
        try:
            # OPTIMIZATION: Cache embedding for general queries (always the same)
            global _general_query_embedding_cache
            if use_general_query and _general_query_embedding_cache is not None:
                query_embedding = _general_query_embedding_cache
            else:
                query_embedding = generate_embeddings(query_text)
                if use_general_query:
                    _general_query_embedding_cache = query_embedding
            
            # Get user's processed materials AND system materials (accessible to all users)
            # If no case study/topic, prioritize system materials
            if use_general_query:
                # Get only system materials for general question generation
                user_materials = db.query(Material).filter(
                    Material.user_id == SYSTEM_USER_ID,
                    Material.status == "processed"
                ).all()
            else:
                # Get both user and system materials
                user_materials = db.query(Material).filter(
                    or_(
                        Material.user_id == current_user['user_id'],
                        Material.user_id == SYSTEM_USER_ID
                    ),
                    Material.status == "processed"
                ).all()
            
            if user_materials:
                material_ids = [m.id for m in user_materials]
                
                # AGGRESSIVE OPTIMIZATION: Drastically limit vector entries for speed
                # For general queries, sample from different materials for diversity
                # For specific queries, get top matches
                if use_general_query:
                    # Limit to 50 entries for very fast processing
                    vector_entries = db.query(VectorIndexEntry).filter(
                        VectorIndexEntry.source_id.in_(material_ids),
                        VectorIndexEntry.source_type == "material",
                        VectorIndexEntry.user_id == SYSTEM_USER_ID
                    ).limit(50).all()
                else:
                    # For specific queries, limit to 30 for speed
                    vector_entries = db.query(VectorIndexEntry).filter(
                        VectorIndexEntry.source_id.in_(material_ids),
                        VectorIndexEntry.source_type == "material",
                        or_(
                            VectorIndexEntry.user_id == current_user['user_id'],
                            VectorIndexEntry.user_id == SYSTEM_USER_ID
                        )
                    ).limit(30).all()
                
                # OPTIMIZATION: Use numpy for faster similarity calculation if available
                try:
                    import numpy as np
                    query_array = np.array(query_embedding)
                    use_numpy = True
                except ImportError:
                    use_numpy = False
                
                # Calculate similarity scores and get top results
                results = []
                for entry in vector_entries:
                    if use_numpy:
                        entry_array = np.array(entry.embedding)
                        # Cosine similarity: dot product / (norm1 * norm2)
                        similarity = np.dot(query_array, entry_array) / (
                            np.linalg.norm(query_array) * np.linalg.norm(entry_array)
                        )
                    else:
                        # Fallback to simple dot product (faster than cosine for ranking)
                        similarity = sum(a * b for a, b in zip(query_embedding, entry.embedding))
                    
                    results.append({
                        "content": entry.content,
                        "similarity": similarity,
                        "metadata": entry.vector_metadata,
                        "is_system": entry.user_id == SYSTEM_USER_ID
                    })
                
                # AGGRESSIVE OPTIMIZATION: Minimal chunks and content for fastest response
                num_results = 3 if use_general_query else 3  # Only 3 chunks max
                results.sort(key=lambda x: x["similarity"], reverse=True)
                top_results = results[:num_results]
                
                if top_results:
                    relevant_context = "\n\nRelevant information:\n"
                    for i, result in enumerate(top_results, 1):
                        # AGGRESSIVE: Very short content snippets (400 chars max)
                        content_length = 400
                        relevant_context += f"\n{result['content'][:content_length]}\n"
        
        except Exception as e:
            error_msg = str(e)
            print(f"RAG search error: {error_msg}")
            # Continue without RAG context if search fails
        
        # AGGRESSIVE OPTIMIZATION: Shorter prompts for faster processing
        if use_general_query:
            system_prompt = f"""Generate {request.num_questions} medical multiple-choice questions. Return ONLY JSON array:
[{{"id":"q1","type":"single","text":"Question?","options":["A","B","C","D"],"correctIndex":0,"explanation":"Why correct"}}]"""
            
            user_prompt = f"""Based on: {relevant_context}\n\nGenerate {request.num_questions} diverse medical questions."""
        else:
            system_prompt = f"""Generate {request.num_questions} medical questions. Return ONLY JSON:
[{{"id":"q1","type":"single","text":"Question?","options":["A","B","C","D"],"correctIndex":0,"explanation":"Why"}}]"""

            user_prompt = f"""{query_text}\n\n{relevant_context}\n\nGenerate {request.num_questions} questions."""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        # AGGRESSIVE OPTIMIZATION: Use fastest model with minimal tokens
        response = client.chat.completions.create(
            model="gpt-4o-mini",  # Fastest and cheapest model
            messages=messages,
            max_tokens=800,  # Minimal tokens for fastest response
            temperature=0.7,
            stream=False
        )
        
        ai_response = response.choices[0].message.content.strip()
        
        # Parse JSON response (handle markdown code blocks if present)
        if ai_response.startswith("```json"):
            ai_response = ai_response[7:]
        if ai_response.startswith("```"):
            ai_response = ai_response[3:]
        if ai_response.endswith("```"):
            ai_response = ai_response[:-3]
        ai_response = ai_response.strip()
        
        try:
            questions = json.loads(ai_response)
            
            # Validate and ensure questions have correct format
            validated_questions = []
            for i, q in enumerate(questions):
                if not isinstance(q, dict):
                    continue
                
                validated_q = {
                    "id": q.get("id", f"q{i+1}"),
                    "type": q.get("type", "single"),
                    "text": q.get("text", ""),
                    "options": q.get("options", []),
                    "correctIndex": q.get("correctIndex", 0),
                    "explanation": q.get("explanation", "")
                }
                
                # Ensure we have exactly 4 options
                if len(validated_q["options"]) != 4:
                    continue
                
                # Ensure correctIndex is valid
                if validated_q["correctIndex"] < 0 or validated_q["correctIndex"] >= 4:
                    validated_q["correctIndex"] = 0
                
                validated_questions.append(validated_q)
            
            if not validated_questions:
                raise ValueError("No valid questions generated")
            
            return {
                "success": True,
                "questions": validated_questions,
                "num_generated": len(validated_questions)
            }
            
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {e}")
            print(f"AI response: {ai_response}")
            raise HTTPException(
                status_code=500,
                detail=f"Error parsing generated questions. AI response: {ai_response[:200]}"
            )
        except Exception as e:
            print(f"Error validating questions: {e}")
            raise HTTPException(status_code=500, detail=f"Error validating generated questions: {str(e)}")
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error generating questions: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating questions: {str(e)}")

# Learning Plan Management Endpoints
class CreateLearningPlanRequest(BaseModel):
    title: str
    description: Optional[str] = None
    video_url: Optional[str] = None
    video_title: Optional[str] = None
    video_duration: Optional[int] = None
    case_study: str
    case_study_editable: bool = False
    quiz_questions: List[Dict[str, Any]]
    topic: Optional[str] = None
    difficulty_level: Optional[str] = None
    estimated_duration: Optional[int] = None

@app.post("/api/learning-plans")
async def create_or_get_learning_plan(
    request: CreateLearningPlanRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new learning plan row for each generated question set"""
    try:
        user_id = current_user['user_id']
        
        # IMPORTANT: Always create a new learning plan row for each generated question set
        # This ensures each set of questions is stored separately, similar to learning_plan_progress
        # Each row will have a unique ID, so they can be distinguished even with the same title
        
        # Create new learning plan row
        learning_plan = LearningPlan(
            user_id=user_id,
            title=request.title,
            description=request.description,
            video_url=request.video_url,
            video_title=request.video_title,
            video_duration=request.video_duration,
            case_study=request.case_study,
            case_study_editable=request.case_study_editable,
            quiz_questions=request.quiz_questions,
            topic=request.topic,
            difficulty_level=request.difficulty_level,
            estimated_duration=request.estimated_duration,
            created_by=user_id
        )
        
        db.add(learning_plan)
        db.commit()
        db.refresh(learning_plan)
        
        return {
            "success": True,
            "learning_plan_id": learning_plan.id,
            "message": "Learning plan created successfully",
            "is_new": True
        }
        
    except Exception as e:
        db.rollback()
        print(f"Error creating learning plan: {e}")
        raise HTTPException(status_code=500, detail=f"Error creating learning plan: {str(e)}")

class SubmitQuizRequest(BaseModel):
    learning_plan_id: int
    quiz_answers: Dict[str, int]  # {questionId: answerIndex}
    video_watched: Optional[bool] = False
    case_study_read: Optional[bool] = False

@app.post("/api/learning-plans/submit-quiz")
async def submit_learning_plan_quiz(
    request: SubmitQuizRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Submit quiz answers and save progress"""
    try:
        user_id = current_user['user_id']
        
        # Verify learning plan exists and belongs to user
        learning_plan = db.query(LearningPlan).filter(
            LearningPlan.id == request.learning_plan_id,
            LearningPlan.user_id == user_id
        ).first()
        
        if not learning_plan:
            raise HTTPException(status_code=404, detail="Learning plan not found")
        
        # Calculate quiz score first
        questions = learning_plan.quiz_questions
        if not isinstance(questions, list):
            questions = []
        
        correct = 0
        total = len(questions)
        
        for question in questions:
            question_id = question.get('id')
            correct_index = question.get('correctIndex')
            # Handle both string and int question IDs
            user_answer = request.quiz_answers.get(str(question_id)) or request.quiz_answers.get(question_id)
            
            if user_answer is not None and user_answer == correct_index:
                correct += 1
        
        percentage = (correct / total * 100) if total > 0 else 0
        
        # Calculate attempt count based on number of existing quiz submissions for this learning plan
        attempt_count = db.query(LearningPlanProgress).filter(
            LearningPlanProgress.user_id == user_id,
            LearningPlanProgress.learning_plan_id == request.learning_plan_id,
            LearningPlanProgress.quiz_submitted == True
        ).count() + 1
        
        # Get or create a base progress record for tracking video/case study (non-quiz progress)
        # We'll use the most recent one or create new if none exists
        base_progress = db.query(LearningPlanProgress).filter(
            LearningPlanProgress.user_id == user_id,
            LearningPlanProgress.learning_plan_id == request.learning_plan_id
        ).order_by(LearningPlanProgress.started_at.desc()).first()
        
        # Create a NEW row for this quiz submission
        # This ensures each generated set of questions gets its own row
        progress = LearningPlanProgress(
            user_id=user_id,
            learning_plan_id=request.learning_plan_id,
            # Copy video/case study status from base progress if it exists
            video_watched=base_progress.video_watched if base_progress else (request.video_watched or False),
            video_watched_at=base_progress.video_watched_at if base_progress and base_progress.video_watched else None,
            video_progress=base_progress.video_progress if base_progress else 0,
            case_study_read=base_progress.case_study_read if base_progress else (request.case_study_read or False),
            case_study_read_at=base_progress.case_study_read_at if base_progress and base_progress.case_study_read else None,
            # Quiz-specific data for this submission
            quiz_submitted=True,
            quiz_submitted_at=func.now(),
            quiz_answers=request.quiz_answers,  # Store answers directly (one set per row)
            quiz_score=correct,
            quiz_total=total,
            quiz_percentage=percentage,
            attempt_count=attempt_count,
            last_accessed=func.now()
        )
        
        # Update video and case study status if provided
        if request.video_watched and not progress.video_watched:
            progress.video_watched = True
            progress.video_watched_at = func.now()
        
        if request.case_study_read and not progress.case_study_read:
            progress.case_study_read = True
            progress.case_study_read_at = func.now()
        
        # Mark as completed if all components are done
        if progress.video_watched and progress.case_study_read and progress.quiz_submitted:
            progress.completed = True
            progress.completed_at = func.now()
        
        db.add(progress)
        
        db.commit()
        db.refresh(progress)
        
        return {
            "success": True,
            "message": "Quiz submitted successfully",
            "score": correct,
            "total": total,
            "percentage": float(percentage),
            "completed": progress.completed
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        print(f"Error submitting quiz: {e}")
        raise HTTPException(status_code=500, detail=f"Error submitting quiz: {str(e)}")

@app.get("/api/learning-plans/{learning_plan_id}/progress")
async def get_learning_plan_progress(
    learning_plan_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get user's progress for a specific learning plan - returns all quiz attempts"""
    try:
        user_id = current_user['user_id']
        
        # Get all progress records for this learning plan (each quiz submission is a row)
        all_progress = db.query(LearningPlanProgress).filter(
            LearningPlanProgress.user_id == user_id,
            LearningPlanProgress.learning_plan_id == learning_plan_id
        ).order_by(LearningPlanProgress.quiz_submitted_at.desc(), LearningPlanProgress.started_at.desc()).all()
        
        if not all_progress:
            return {
                "success": True,
                "progress": None,
                "quiz_attempts": [],
                "message": "No progress found"
            }
        
        # Get the most recent progress for video/case study status
        latest_progress = all_progress[0]
        
        # Get all quiz attempts (rows where quiz was submitted)
        quiz_attempts = []
        for prog in all_progress:
            if prog.quiz_submitted:
                quiz_attempts.append({
                    "id": prog.id,
                    "quiz_submitted_at": prog.quiz_submitted_at.isoformat() if prog.quiz_submitted_at else None,
                    "quiz_answers": prog.quiz_answers,
                    "quiz_score": prog.quiz_score,
                    "quiz_total": prog.quiz_total,
                    "quiz_percentage": float(prog.quiz_percentage) if prog.quiz_percentage else None,
                    "attempt_number": prog.attempt_count
                })
        
        return {
            "success": True,
            "progress": {
                "id": latest_progress.id,
                "video_watched": latest_progress.video_watched,
                "video_watched_at": latest_progress.video_watched_at.isoformat() if latest_progress.video_watched_at else None,
                "video_progress": latest_progress.video_progress,
                "case_study_read": latest_progress.case_study_read,
                "case_study_read_at": latest_progress.case_study_read_at.isoformat() if latest_progress.case_study_read_at else None,
                "completed": latest_progress.completed,
                "completed_at": latest_progress.completed_at.isoformat() if latest_progress.completed_at else None,
                "started_at": latest_progress.started_at.isoformat() if latest_progress.started_at else None,
                "total_quiz_attempts": len(quiz_attempts)
            },
            "quiz_attempts": quiz_attempts  # All quiz submissions, each in its own row
        }
        
    except Exception as e:
        print(f"Error fetching progress: {e}")
        raise HTTPException(status_code=500, detail=f"Error fetching progress: {str(e)}")

# Profile API Endpoints
def _get_user_uuid(user_id):
    """Helper function to convert user_id string to UUID"""
    from uuid import UUID as UUIDType
    if isinstance(user_id, str):
        try:
            return UUIDType(user_id)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid user ID format: {user_id}")
    return user_id

@app.post("/api/profile")
async def create_or_update_profile(
    profile_data: Dict[str, Any] = Body(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Create or update user profile"""
    try:
        # Convert user_id string to UUID if needed
        user_id_str = current_user["user_id"]
        print(f"DEBUG: Original user_id: {user_id_str}, type: {type(user_id_str)}")
        
        user_id_uuid = _get_user_uuid(user_id_str)
        print(f"DEBUG: Converted user_id UUID: {user_id_uuid}, type: {type(user_id_uuid)}")
        print(f"DEBUG: Profile data: {profile_data}")
        
        # Check if profile exists with string ID (in case of mismatch)
        existing_profile_str = None
        try:
            existing_profile_str = (
                db.query(Profile)
                .filter(Profile.id == user_id_str)
                .first()
            )
            if existing_profile_str:
                print(f"DEBUG: Found profile with string ID: {existing_profile_str.id}")
        except Exception as e:
            print(f"DEBUG: Error checking string ID (expected if UUID type): {e}")
        
        # clean grad_year so "" doesn't break integer column
        raw_grad_year = profile_data.get("grad_year")
        if isinstance(raw_grad_year, str):
            raw_grad_year = raw_grad_year.strip() or None

        # Check if profile already exists with UUID
        existing_profile = (
            db.query(Profile)
            .filter(Profile.id == user_id_uuid)
            .first()
        )
        print(f"DEBUG: Existing profile found (UUID): {existing_profile is not None}")
        
        # If found with string but not UUID, use the string one
        if existing_profile_str and not existing_profile:
            print(f"DEBUG: Using profile found with string ID")
            existing_profile = existing_profile_str

        if existing_profile:
            # Update existing profile
            existing_profile.full_name = profile_data.get("full_name")
            existing_profile.institution = profile_data.get("institution")
            existing_profile.grad_year = raw_grad_year
            existing_profile.specialty = profile_data.get("specialty")

            # only set updated_at if the column actually exists
            if hasattr(existing_profile, "updated_at"):
                existing_profile.updated_at = func.now()

            print(f"DEBUG: Updating existing profile, committing...")
            db.commit()
            print(f"DEBUG: Update commit successful, refreshing...")
            db.refresh(existing_profile)
            print(f"DEBUG: Profile updated: {existing_profile.id}, {existing_profile.full_name}")

            return {
                "success": True,
                "message": "Profile updated successfully",
                "profile": {
                    "id": existing_profile.id,
                    "full_name": existing_profile.full_name,
                    "institution": existing_profile.institution,
                    "grad_year": existing_profile.grad_year,
                    "specialty": existing_profile.specialty,
                    # guard updated_at in case it doesn't exist / is None
                    "updated_at": (
                        existing_profile.updated_at.isoformat()
                        if getattr(existing_profile, "updated_at", None)
                        else None
                    ),
                },
            }
        else:
            # Create new profile
            # First, check if there's already a profile with this ID in any format
            print(f"DEBUG: No existing profile found, creating new one...")
            
            # List all existing profiles to debug
            try:
                all_profiles = db.query(Profile).limit(10).all()
                print(f"DEBUG: Sample of existing profiles in DB ({len(all_profiles)} shown):")
                for p in all_profiles:
                    print(f"  - Profile ID: {p.id} (type: {type(p.id).__name__}), name: {p.full_name}")
            except Exception as e:
                print(f"DEBUG: Could not list profiles: {e}")
            
            try:
                profile = Profile(
                    id=user_id_uuid,
                    full_name=profile_data.get("full_name"),
                    institution=profile_data.get("institution"),
                    grad_year=raw_grad_year,
                    specialty=profile_data.get("specialty"),
                )

                db.add(profile)
                print(f"DEBUG: Profile added to session, attempting commit...")
                db.commit()
                print(f"DEBUG: Commit successful, refreshing profile...")
                db.refresh(profile)
                print(f"DEBUG: Profile refreshed: {profile.id}, {profile.full_name}")
            except Exception as create_error:
                db.rollback()
                error_type = type(create_error).__name__
                error_msg = str(create_error)
                print(f"DEBUG: Error creating profile!")
                print(f"DEBUG: Error type: {error_type}")
                print(f"DEBUG: Error message: {error_msg}")
                
                # Check for common issues
                if "duplicate" in error_msg.lower() or "unique" in error_msg.lower():
                    print(f"DEBUG: Possible duplicate key error - profile may already exist")
                    # Try to find it
                    try:
                        found = db.query(Profile).filter(Profile.id == user_id_str).first()
                        if found:
                            print(f"DEBUG: Found existing profile with string ID: {found.id}")
                        found_uuid = db.query(Profile).filter(Profile.id == user_id_uuid).first()
                        if found_uuid:
                            print(f"DEBUG: Found existing profile with UUID: {found_uuid.id}")
                    except:
                        pass
                
                # Re-raise to show the actual error
                raise

            return {
                "success": True,
                "message": "Profile created successfully",
                "profile": {
                    "id": profile.id,
                    "full_name": profile.full_name,
                    "institution": profile.institution,
                    "grad_year": profile.grad_year,
                    "specialty": profile.specialty,
                    "created_at": (
                        profile.created_at.isoformat()
                        if getattr(profile, "created_at", None)
                        else None
                    ),
                },
            }

    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        db.rollback()
        # also log this on the server so you can see the exact DB error
        import traceback
        error_trace = traceback.format_exc()
        print("=" * 50)
        print("ERROR SAVING PROFILE:")
        print(f"Error type: {type(e).__name__}")
        print(f"Error message: {str(e)}")
        print("Full traceback:")
        print(error_trace)
        print("=" * 50)
        raise HTTPException(status_code=500, detail=f"Error saving profile: {str(e)}")

@app.get("/api/profile")
async def get_profile(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get user profile"""
    try:
        # Convert user_id string to UUID if needed
        user_id_uuid = _get_user_uuid(current_user["user_id"])
        
        profile = (
            db.query(Profile)
            .filter(Profile.id == user_id_uuid)
            .first()
        )

        if not profile:
            return {
                "success": True,
                "profile": None,
                "message": "No profile found",
            }

        return {
            "success": True,
            "profile": {
                "id": profile.id,
                "full_name": profile.full_name,
                "institution": profile.institution,
                "grad_year": profile.grad_year,
                "specialty": profile.specialty,
                "created_at": (
                    profile.created_at.isoformat()
                    if getattr(profile, "created_at", None)
                    else None
                ),
                "updated_at": (
                    profile.updated_at.isoformat()
                    if getattr(profile, "updated_at", None)
                    else None
                ),
            },
        }

    except Exception as e:
        print("Error fetching profile:", e)
        raise HTTPException(
            status_code=500, detail=f"Error fetching profile: {str(e)}"
        )


@app.get("/api/profile/me")
async def get_my_profile_with_user_data(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get combined user data from Supabase and profile from database"""
    try:
        # Convert user_id string to UUID if needed
        user_id_uuid = _get_user_uuid(current_user["user_id"])
        
        profile = (
            db.query(Profile)
            .filter(Profile.id == user_id_uuid)
            .first()
        )

        user_data = {
            "user_id": current_user["user_id"],
            "email": current_user.get("email"),
            "full_name": profile.full_name if profile else None,
            "institution": profile.institution if profile else None,
            "grad_year": profile.grad_year if profile else None,
            "specialty": profile.specialty if profile else None,
            "profile_exists": profile is not None,
        }

        return {
            "success": True,
            "profile": user_data,
        }

    except Exception as e:
        print("Error fetching user data:", e)
        raise HTTPException(
            status_code=500, detail=f"Error fetching user data: {str(e)}"
        )


# Helper functions for care plan operations
def build_care_plan_text(care_plan_data: dict) -> str:
    """Build text representation of care plan for RAG indexing"""
    sections = []
    
    # Patient Information
    if care_plan_data.get('patient_name'):
        sections.append(f"Patient: {care_plan_data['patient_name']}")
    if care_plan_data.get('age'):
        sections.append(f"Age: {care_plan_data['age']}")
    if care_plan_data.get('sex'):
        sections.append(f"Sex: {care_plan_data['sex']}")
    if care_plan_data.get('diagnosis'):
        sections.append(f"Diagnosis: {care_plan_data['diagnosis']}")
    if care_plan_data.get('procedure'):
        sections.append(f"Procedure: {care_plan_data['procedure']}")
    
    # Medical History
    if care_plan_data.get('pmh'):
        sections.append(f"Past Medical History: {care_plan_data['pmh']}")
    if care_plan_data.get('psh'):
        sections.append(f"Past Surgical History: {care_plan_data['psh']}")
    if care_plan_data.get('meds'):
        sections.append(f"Current Medications: {care_plan_data['meds']}")
    
    # Physical Assessment
    if care_plan_data.get('neuro'):
        sections.append(f"Neurological: {care_plan_data['neuro']}")
    if care_plan_data.get('cardio'):
        sections.append(f"Cardiovascular: {care_plan_data['cardio']}")
    if care_plan_data.get('resp'):
        sections.append(f"Respiratory: {care_plan_data['resp']}")
    
    # Airway Assessment
    if care_plan_data.get('mallampati'):
        sections.append(f"Mallampati: {care_plan_data['mallampati']}")
    if care_plan_data.get('ulbt'):
        sections.append(f"ULBT Grade: {care_plan_data['ulbt']}")
    
    # Laboratory Values
    lab_values = []
    for lab in ['na', 'k', 'cl', 'co2', 'bun', 'cr', 'glu', 'wbc', 'hgb', 'hct', 'plt', 'pt', 'ptt', 'inr']:
        if care_plan_data.get(lab):
            lab_values.append(f"{lab.upper()}: {care_plan_data[lab]}")
    if lab_values:
        sections.append(f"Laboratory Values: {', '.join(lab_values)}")
    
    return "\n".join(sections)

async def index_care_plan_for_rag(care_plan: CarePlan, db: Session):
    """Index care plan content for RAG retrieval"""
    try:
        if not care_plan.exported_text:
            return
        
        # Create vector index entry for the care plan
        embedding = generate_embeddings(care_plan.exported_text)
        
        vector_entry = VectorIndexEntry(
            user_id=care_plan.user_id,
            content_hash=f"care_plan_{care_plan.id}",
            embedding=embedding,
            content=care_plan.exported_text,
            token_count=len(care_plan.exported_text.split()),
            chunk_index=0,
            source_type="care_plan",
            source_id=care_plan.id,
            embedding_model="text-embedding-3-small",
            vector_metadata={
                "care_plan_id": care_plan.id,
                "patient_name": care_plan.patient_name,
                "procedure": care_plan.procedure,
                "diagnosis": care_plan.diagnosis,
                "created_at": care_plan.created_at.isoformat() if care_plan.created_at else None
            }
        )
        
        db.add(vector_entry)
        db.commit()
        
    except Exception as e:
        print(f"Error indexing care plan for RAG: {e}")

def build_care_plan_context(care_plan: CarePlan, db: Session) -> str:
    """Build context for AI generation by searching relevant materials"""
    try:
        # Search for relevant materials using care plan content
        query_text = f"{care_plan.diagnosis} {care_plan.procedure} anesthesia care plan"
        query_embedding = generate_embeddings(query_text)
        
        # Get relevant vector entries (user's materials AND system materials)
        vector_entries = db.query(VectorIndexEntry).filter(
            or_(
                VectorIndexEntry.user_id == care_plan.user_id,
                VectorIndexEntry.user_id == SYSTEM_USER_ID
            )
        ).all()
        
        # Calculate similarity and get top results
        results = []
        for entry in vector_entries:
            similarity = sum(a * b for a, b in zip(query_embedding, entry.embedding))
            results.append({
                "content": entry.content,
                "similarity": similarity,
                "metadata": entry.vector_metadata,
                "is_system": entry.user_id == SYSTEM_USER_ID
            })
        
        # Sort by similarity and get top 5 most relevant chunks
        results.sort(key=lambda x: x["similarity"], reverse=True)
        top_results = results[:5]
        
        context_parts = [f"Patient Care Plan:\n{care_plan.exported_text}"]
        
        if top_results:
            context_parts.append("\nRelevant Medical Literature:")
            for i, result in enumerate(top_results, 1):
                source_name = result["metadata"].get("file_name", "Unknown Source")
                source_label = "System Textbook" if result.get("is_system") else "User Material"
                context_parts.append(f"\n{i}. From {source_name} ({source_label}):\n{result['content'][:300]}...")
        
        return "\n".join(context_parts)
        
    except Exception as e:
        print(f"Error building care plan context: {e}")
        return care_plan.exported_text or ""

async def generate_care_plan_recommendations(context: str, client) -> dict:
    """Generate AI recommendations for care plan"""
    try:
        system_prompt = """You are an expert anesthesiologist AI assistant. Based on the patient information and relevant medical literature provided, generate comprehensive anesthesia care plan recommendations.

Please provide:
1. Anesthesia Plan: Detailed anesthesia approach including induction, maintenance, and emergence
2. Risk Assessment: Analysis of patient-specific risks and complications
3. Monitoring Plan: Required monitoring during the procedure
4. Medication Plan: Specific medications and dosages

Be specific, evidence-based, and consider the patient's comorbidities and procedure requirements."""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": context}
        ]
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=1500
        )
        
        ai_response = response.choices[0].message.content
        
        # Parse the response into structured recommendations
        recommendations = {
            "anesthesia_plan": ai_response,
            "risk_assessment": "",
            "monitoring_plan": "",
            "medication_plan": "",
            "sources": [],
            "confidence_score": 0.8
        }
        
        return recommendations
        
    except Exception as e:
        print(f"Error generating AI recommendations: {e}")
        return {
            "anesthesia_plan": "Error generating recommendations",
            "risk_assessment": "",
            "monitoring_plan": "",
            "medication_plan": "",
            "sources": [],
            "confidence_score": 0.0
        }
